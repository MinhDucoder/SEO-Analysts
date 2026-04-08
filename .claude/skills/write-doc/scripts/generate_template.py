#!/usr/bin/env python3
"""
Generate a DOCX template with correct formatting for Vietnamese academic reports.

Usage:
    python generate_template.py output.docx
    python generate_template.py output.docx --title "Tên đồ án"
    python generate_template.py output.docx --title "Tên đồ án" --skeleton
"""

import argparse
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# === Constants from formatting rules ===
A4_WIDTH = Inches(8.27)
A4_HEIGHT = Inches(11.69)
MARGIN = Inches(1)
HEADER_FOOTER_DIST = Inches(0.5)

HEADING_COLOR = RGBColor(0x0F, 0x47, 0x61)
BODY_COLOR = RGBColor(0x00, 0x00, 0x00)

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri Light"

SIZE_BODY = Pt(12)
SIZE_H1 = Pt(20)
SIZE_H2 = Pt(16)
SIZE_H3 = Pt(14)

LINE_SPACING = 1.15
SPACE_AFTER_BODY = Pt(8)

H1_BEFORE = Pt(18)
H1_AFTER = Pt(4)
H2_BEFORE = Pt(8)
H2_AFTER = Pt(4)
H3_BEFORE = Pt(8)
H3_AFTER = Pt(4)


def setup_page(doc):
    """Configure A4 page with 1-inch margins."""
    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = HEADER_FOOTER_DIST
    section.footer_distance = HEADER_FOOTER_DIST


def set_font_vietnamese(run, font_name, size, color, bold=False):
    """Set font with vi-VN language support."""
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    # Ensure East Asian and complex script fonts are also set
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    # Set language to vi-VN
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = r.makeelement(qn('w:lang'), {})
        rPr.append(lang)
    lang.set(qn('w:val'), 'vi-VN')


def setup_styles(doc):
    """Configure Normal and Heading 1/2/3 styles."""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = SIZE_BODY
    font.color.rgb = BODY_COLOR
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = SPACE_AFTER_BODY
    pf.space_before = Pt(0)

    # Heading styles
    heading_configs = [
        ('Heading 1', SIZE_H1, H1_BEFORE, H1_AFTER),
        ('Heading 2', SIZE_H2, H2_BEFORE, H2_AFTER),
        ('Heading 3', SIZE_H3, H3_BEFORE, H3_AFTER),
    ]

    for style_name, size, before, after in heading_configs:
        style = doc.styles[style_name]
        font = style.font
        font.name = FONT_HEADING
        font.size = size
        font.color.rgb = HEADING_COLOR
        font.bold = True
        pf = style.paragraph_format
        pf.space_before = before
        pf.space_after = after
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = LINE_SPACING


# === Document structure (7 parts) ===
SKELETON = {
    "chapters": [
        {
            "title": "Chương 1: Tổng quan về bài toán, công nghệ, công cụ",
            "sections": [
                ("1.1 Hiện trạng", [
                    "Chi tiết vấn đề cần giải quyết",
                    "Phần mềm/hệ thống hiện tại",
                    "Vướng mắc, hạn chế → yêu cầu bài toán",
                ]),
                ("1.2 Bài toán", [
                    "Yêu cầu hệ thống tổng thể",
                    "Yêu cầu chức năng",
                    "Yêu cầu phi chức năng",
                    "Đối tượng người dùng",
                ]),
                ("1.3 Công nghệ và công cụ", [
                    "Công cụ phân tích thiết kế",
                    "Công nghệ lập trình (ngôn ngữ, framework)",
                    "Công cụ quản lý dự án",
                    "Cơ sở dữ liệu",
                    "Công cụ đồ họa (nếu có)",
                    "[Lưu ý: Mỗi công nghệ PHẢI dẫn nguồn tài liệu tham khảo]",
                ]),
                ("1.4 Quy trình phát triển phần mềm", [
                    "Mô hình phát triển áp dụng",
                    "Lý do chọn mô hình",
                ]),
            ],
        },
        {
            "title": "Chương 2: Phân tích và thiết kế hệ thống",
            "sections": [
                ("2.1 Hướng cấu trúc (hoặc hướng đối tượng)", [
                    "BFD / Use Case Diagram",
                    "DFD / Sequence Diagram",
                    "ERD / Class Diagram",
                    "[Lưu ý: ERD ≠ Class Diagram. Mọi biểu đồ PHẢI kèm phân tích giải thích]",
                ]),
            ],
        },
        {
            "title": "Chương 3: Xây dựng ứng dụng",
            "sections": [
                ("3.1 Kiến trúc hệ thống", [
                    "Sơ đồ kiến trúc tổng thể",
                    "Giải thích các tầng/lớp",
                ]),
                ("3.2 Triển khai xây dựng", [
                    "Quản lý source code",
                    "Quản lý tiến độ (kèm screenshot Trello/Jira)",
                ]),
                ("3.3 Một số kết quả", []),
                ("3.3.1 Giao diện và chức năng dành cho quản lý (admin)", [
                    "Screenshot + mô tả chức năng",
                ]),
                ("3.3.2 Giao diện và chức năng dành cho khách", [
                    "Screenshot + mô tả chức năng",
                ]),
            ],
        },
        {
            "title": "Chương 4: Kiểm thử và triển khai sản phẩm",
            "sections": [
                ("4.1 Kiểm thử", [
                    "Chiến lược kiểm thử",
                    "Thiết kế test case (bảng test case)",
                    "Kết quả kiểm thử (pass/fail + evidence)",
                ]),
                ("4.2 Triển khai dự án (nếu có)", [
                    "Host ở đâu",
                    "Triển khai như thế nào",
                ]),
            ],
        },
    ],
    "conclusion": [
        "Kiến thức, kỹ năng đạt được so với mục tiêu ban đầu",
        "Ưu điểm / nhược điểm của sản phẩm",
        "Hướng phát triển tương lai",
        "Chi tiết kỹ năng mềm (làm việc nhóm, sử dụng công cụ...)",
    ],
    "references_note": (
        "Sắp xếp theo thể loại: Sách → Website/Báo → Tài liệu kỹ thuật. "
        "Trong mỗi nhóm xếp ABC theo tên tác giả. "
        "Website PHẢI ghi ngày truy cập. 15-20 tài liệu."
    ),
    "appendix_note": (
        "Đánh số trang riêng (PL-1, PL-2... hoặc i, ii, iii...). >= 30 trang. "
        "Bao gồm: Bảng phân công thành viên, Khảo sát, Test case chi tiết."
    ),
}


def add_placeholder(doc, text):
    """Add a placeholder paragraph in italic gray."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = SIZE_BODY
    run.font.name = FONT_BODY


def build_skeleton(doc, title=None):
    """Build document with full 7-part structure and placeholders."""
    # Title page placeholder
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_font_vietnamese(run, FONT_HEADING, Pt(26), HEADING_COLOR, bold=True)
        p.paragraph_format.space_before = Pt(200)
        p.paragraph_format.space_after = Pt(40)
        doc.add_page_break()

    # Chapters
    for ch in SKELETON["chapters"]:
        doc.add_heading(ch["title"], level=1)
        for section_title, items in ch["sections"]:
            # Use H2 for main sections, H3 for sub-sections (e.g., 3.3.1)
            level = 3 if "." in section_title.split(" ")[0] and section_title.split(" ")[0].count(".") >= 2 else 2
            doc.add_heading(section_title, level=level)
            for item in items:
                if item.startswith("["):
                    add_placeholder(doc, item[1:-1])
                else:
                    add_placeholder(doc, item)

    # Conclusion (NOT a chapter - continuous paragraphs)
    doc.add_page_break()
    p = doc.add_heading("Kết luận", level=1)
    add_placeholder(doc, "KHÔNG chia tiểu mục — viết thành các đoạn văn liên tục")
    for item in SKELETON["conclusion"]:
        add_placeholder(doc, item)

    # References
    doc.add_page_break()
    doc.add_heading("Tài liệu tham khảo", level=1)
    add_placeholder(doc, SKELETON["references_note"])

    # Appendix
    doc.add_page_break()
    doc.add_heading("Phụ lục", level=1)
    add_placeholder(doc, SKELETON["appendix_note"])


def generate(output_path, title=None, skeleton=False):
    """Generate the DOCX template."""
    doc = Document()
    setup_page(doc)
    setup_styles(doc)

    if skeleton:
        build_skeleton(doc, title)
    else:
        # Minimal template with just styles configured
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            set_font_vietnamese(run, FONT_HEADING, Pt(26), HEADING_COLOR, bold=True)

        doc.add_heading("Chương 1: Tổng quan", level=1)
        doc.add_paragraph("Nội dung chương 1...")
        doc.add_heading("1.1 Hiện trạng", level=2)
        doc.add_paragraph("Nội dung mục 1.1...")

    doc.save(output_path)
    print(f"OK: {output_path}")
    if skeleton:
        print("  Mode: skeleton (full 7-part structure with placeholders)")
    else:
        print("  Mode: minimal (styles configured, sample headings)")


def main():
    parser = argparse.ArgumentParser(
        description="Generate DOCX template for Vietnamese academic reports"
    )
    parser.add_argument("output", help="Output .docx file path")
    parser.add_argument("--title", help="Report title", default=None)
    parser.add_argument(
        "--skeleton",
        action="store_true",
        help="Generate full 7-part structure with placeholders",
    )
    args = parser.parse_args()

    if not args.output.endswith(".docx"):
        args.output += ".docx"

    generate(args.output, args.title, args.skeleton)


if __name__ == "__main__":
    main()

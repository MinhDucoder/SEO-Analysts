#!/usr/bin/env python3
"""
Validate a DOCX file against Vietnamese academic report formatting rules.

Usage:
    python validate_docx.py report.docx
    python validate_docx.py report.docx --strict    # Also check structure
"""

import argparse
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Emu
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# === Expected values ===
EXPECTED = {
    "page_width_emu": Inches(8.27),
    "page_height_emu": Inches(11.69),
    "margin_emu": Inches(1),
    "margin_tolerance_emu": Inches(0.05),
    "page_tolerance_emu": Inches(0.15),
    "body_font": "Calibri",
    "heading_font": "Calibri Light",
    "body_size_pt": 12,
    "h1_size_pt": 20,
    "h2_size_pt": 16,
    "h3_size_pt": 14,
    "heading_color_hex": "0F4761",
    "line_spacing": 1.15,
    "line_spacing_tolerance": 0.05,
}

REQUIRED_HEADINGS_KEYWORDS = [
    "chương 1",
    "chương 2",
    "chương 3",
    "chương 4",
    "kết luận",
    "tài liệu tham khảo",
    "phụ lục",
]


class ValidationResult:
    def __init__(self):
        self.passed = []
        self.failed = []
        self.warnings = []

    def ok(self, msg):
        self.passed.append(msg)

    def fail(self, msg):
        self.failed.append(msg)

    def warn(self, msg):
        self.warnings.append(msg)

    def report(self):
        total = len(self.passed) + len(self.failed)
        print(f"\n{'='*50}")
        print(f"VALIDATION RESULT: {len(self.passed)}/{total} passed")
        print(f"{'='*50}\n")

        if self.passed:
            print("PASSED:")
            for p in self.passed:
                print(f"  [OK] {p}")

        if self.warnings:
            print(f"\nWARNINGS ({len(self.warnings)}):")
            for w in self.warnings:
                print(f"  [!!] {w}")

        if self.failed:
            print(f"\nFAILED ({len(self.failed)}):")
            for f in self.failed:
                print(f"  [FAIL] {f}")

        print()
        return len(self.failed) == 0


def emu_to_inches(emu):
    return round(emu / 914400, 2)


def validate_page_setup(doc, result):
    """Check page size and margins."""
    section = doc.sections[0]
    tol = EXPECTED["margin_tolerance_emu"]
    page_tol = EXPECTED["page_tolerance_emu"]

    # Page size
    w = section.page_width
    h = section.page_height
    exp_w = EXPECTED["page_width_emu"]
    exp_h = EXPECTED["page_height_emu"]

    if abs(w - exp_w) <= page_tol and abs(h - exp_h) <= page_tol:
        result.ok(f"Page size: A4 ({emu_to_inches(w)}x{emu_to_inches(h)} inch)")
    else:
        result.fail(
            f"Page size: got {emu_to_inches(w)}x{emu_to_inches(h)} inch, "
            f"expected ~8.27x11.69 inch (A4)"
        )

    # Margins
    margins = {
        "top": section.top_margin,
        "bottom": section.bottom_margin,
        "left": section.left_margin,
        "right": section.right_margin,
    }
    exp_margin = EXPECTED["margin_emu"]
    all_ok = True
    for name, val in margins.items():
        if val is None or abs(val - exp_margin) > tol:
            actual = emu_to_inches(val) if val else "None"
            result.fail(f"Margin {name}: got {actual} inch, expected 1.0 inch")
            all_ok = False
    if all_ok:
        result.ok("Margins: 1 inch all sides")


def validate_styles(doc, result):
    """Check font styles for Normal and Headings."""
    # Normal style
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        result.warn("Normal style not found")
        normal = None
    if normal:
        font = normal.font
        if font.name and "Calibri" in font.name and "Light" not in font.name:
            result.ok(f"Normal font: {font.name}")
        elif font.name:
            result.fail(f"Normal font: got '{font.name}', expected 'Calibri'")
        else:
            result.warn("Normal font: not explicitly set (inherits default)")

        if font.size and abs(font.size - Pt(12)) < Pt(0.5):
            result.ok(f"Normal size: {font.size.pt}pt")
        elif font.size:
            result.fail(f"Normal size: got {font.size.pt}pt, expected 12pt")
        else:
            result.warn("Normal size: not explicitly set")

    # Heading styles
    heading_checks = [
        ("Heading 1", EXPECTED["heading_font"], EXPECTED["h1_size_pt"]),
        ("Heading 2", EXPECTED["heading_font"], EXPECTED["h2_size_pt"]),
        ("Heading 3", EXPECTED["heading_font"], EXPECTED["h3_size_pt"]),
    ]

    for style_name, exp_font, exp_size in heading_checks:
        try:
            style = doc.styles[style_name]
        except KeyError:
            result.warn(f"{style_name}: style not found")
            continue

        font = style.font
        # Font name
        if font.name and exp_font.lower() in font.name.lower():
            result.ok(f"{style_name} font: {font.name}")
        elif font.name:
            result.fail(f"{style_name} font: got '{font.name}', expected '{exp_font}'")
        else:
            result.warn(f"{style_name} font: not explicitly set")

        # Font size
        if font.size and abs(font.size - Pt(exp_size)) < Pt(0.5):
            result.ok(f"{style_name} size: {font.size.pt}pt")
        elif font.size:
            result.fail(f"{style_name} size: got {font.size.pt}pt, expected {exp_size}pt")
        else:
            result.warn(f"{style_name} size: not explicitly set")

        # Color
        if font.color and font.color.rgb:
            color_hex = str(font.color.rgb).upper()
            if color_hex == EXPECTED["heading_color_hex"].upper():
                result.ok(f"{style_name} color: #{color_hex}")
            else:
                result.fail(
                    f"{style_name} color: got #{color_hex}, "
                    f"expected #{EXPECTED['heading_color_hex']}"
                )
        else:
            result.warn(f"{style_name} color: not explicitly set")


def validate_spacing(doc, result):
    """Check line spacing and paragraph spacing."""
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        normal = None
    if not normal:
        result.warn("Cannot check spacing: Normal style not found")
        return

    pf = normal.paragraph_format

    # Line spacing
    if pf.line_spacing:
        tol = EXPECTED["line_spacing_tolerance"]
        if abs(pf.line_spacing - EXPECTED["line_spacing"]) <= tol:
            result.ok(f"Line spacing: {pf.line_spacing}")
        else:
            result.fail(
                f"Line spacing: got {pf.line_spacing}, expected ~{EXPECTED['line_spacing']}"
            )
    else:
        result.warn("Line spacing: not explicitly set")

    # Space after
    if pf.space_after:
        if abs(pf.space_after - Pt(8)) < Pt(1):
            result.ok(f"Space after paragraph: {pf.space_after.pt}pt")
        else:
            result.fail(f"Space after: got {pf.space_after.pt}pt, expected 8pt")
    else:
        result.warn("Space after: not explicitly set")


def validate_structure(doc, result):
    """Check document has required 7-part structure (--strict mode)."""
    headings_found = []
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headings_found.append(para.text.lower().strip())

    for keyword in REQUIRED_HEADINGS_KEYWORDS:
        found = any(keyword in h for h in headings_found)
        if found:
            result.ok(f"Structure: '{keyword}' found")
        else:
            result.fail(f"Structure: '{keyword}' NOT found in headings")

    # Check conclusion is not split into sub-headings
    in_conclusion = False
    conclusion_sub_headings = 0
    for para in doc.paragraphs:
        if para.style and para.style.name:
            text_lower = para.text.lower().strip()
            if "kết luận" in text_lower and para.style.name.startswith("Heading"):
                in_conclusion = True
                continue
            if in_conclusion:
                if para.style.name == "Heading 1":
                    break  # Next chapter
                if para.style.name in ("Heading 2", "Heading 3"):
                    conclusion_sub_headings += 1

    if conclusion_sub_headings == 0:
        result.ok("Conclusion: no sub-headings (correct)")
    else:
        result.fail(
            f"Conclusion: found {conclusion_sub_headings} sub-heading(s). "
            "Kết luận KHÔNG được chia tiểu mục"
        )


def main():
    parser = argparse.ArgumentParser(
        description="Validate DOCX against Vietnamese academic report rules"
    )
    parser.add_argument("input", help="Input .docx file to validate")
    parser.add_argument(
        "--strict",
        action="store_true",
        help="Also check 7-part document structure",
    )
    args = parser.parse_args()

    try:
        doc = Document(args.input)
    except Exception as e:
        print(f"Error opening {args.input}: {e}")
        sys.exit(1)

    result = ValidationResult()

    validate_page_setup(doc, result)
    validate_styles(doc, result)
    validate_spacing(doc, result)

    if args.strict:
        validate_structure(doc, result)

    success = result.report()
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()

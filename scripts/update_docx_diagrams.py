"""
Script to update SEO_Report_v2.docx with new microservices diagrams.

Actions:
1. REPLACE 3 existing images (ERD, Audit sequence, Architecture)
2. ADD 3 new diagram sections (Docker deployment, Crawler flow, Rule engine)
3. UPDATE captions and descriptions
4. UPDATE DANH SÁCH HÌNH (figure list)

Output: SEO_Report_v3.docx
"""

import copy
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE_DIR = Path(__file__).parent.parent
DIAGRAMS_DIR = BASE_DIR / "diagrams"
INPUT_DOCX = BASE_DIR / "SEO_Report_v2.docx"
OUTPUT_DOCX = BASE_DIR / "SEO_Report_v3.docx"


def replace_image_in_paragraph(doc, para_index, new_image_path, width_inches=6.0):
    """Replace the image in a paragraph with a new one, keeping the same embed ID."""
    para = doc.paragraphs[para_index]

    # Find the blip element to get embed ID
    blips = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    if not blips:
        print(f"  WARNING: No image found in paragraph {para_index}")
        return

    for blip in blips:
        embed_id = blip.get(qn('r:embed'))
        if embed_id:
            # Get the relationship and replace the image data
            rel = doc.part.rels[embed_id]
            with open(new_image_path, 'rb') as f:
                rel.target_part._blob = f.read()
            print(f"  Replaced image at para {para_index} (embed={embed_id}) with {new_image_path.name}")


def update_paragraph_text(doc, para_index, new_text):
    """Update text of a paragraph while preserving formatting of first run."""
    para = doc.paragraphs[para_index]
    if para.runs:
        # Preserve formatting from first run
        first_run = para.runs[0]
        font_props = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
        }
        # Clear all runs
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
    else:
        para.text = new_text
    print(f"  Updated para {para_index} text: {new_text[:80]}...")


def add_image_paragraph(doc, after_element, image_path, width_inches=6.0):
    """Insert a new paragraph with an image after the specified element."""
    # after_element can be a CT_P element or a paragraph index
    if isinstance(after_element, int):
        after_el = doc.paragraphs[after_element]._element
    else:
        after_el = after_element
    p_new = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    after_el.addnext(p_new)

    # Get the new paragraph object
    # Find it in paragraphs list
    new_p = None
    for p in doc.paragraphs:
        if p._element is p_new:
            new_p = p
            break

    if new_p:
        run = new_p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        print(f"  Added image: {image_path.name}")
    return p_new


def add_text_paragraph(doc, after_element, text, style='Normal', bold=False, italic=False, alignment=None):
    """Insert a new text paragraph after the specified element."""
    p_new = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    after_element.addnext(p_new)

    # Find the new paragraph
    new_p = None
    for p in doc.paragraphs:
        if p._element is p_new:
            new_p = p
            break

    if new_p:
        new_p.style = doc.styles[style] if style in [s.name for s in doc.styles] else doc.styles['Normal']
        run = new_p.add_run(text)
        run.bold = bold
        run.italic = italic
        if run.font:
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
        if alignment:
            new_p.alignment = alignment
        print(f"  Added text para: {text[:80]}...")

    return p_new


def add_caption_paragraph(doc, after_element, caption_text):
    """Add a centered, italic caption paragraph."""
    p_new = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    after_element.addnext(p_new)

    new_p = None
    for p in doc.paragraphs:
        if p._element is p_new:
            new_p = p
            break

    if new_p:
        run = new_p.add_run(caption_text)
        run.bold = True
        run.italic = True
        if run.font:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return p_new


def main():
    print("=" * 60)
    print("Updating SEO_Report_v2.docx → SEO_Report_v3.docx")
    print("=" * 60)

    doc = Document(str(INPUT_DOCX))

    # ============================================================
    # STEP 1: Replace existing images
    # ============================================================
    print("\n--- Step 1: Replace existing images ---")

    # Hình 2.1 ERD (para 390) → 14-erd-microservices.png
    replace_image_in_paragraph(doc, 390, DIAGRAMS_DIR / "14-erd-microservices.png")
    update_paragraph_text(doc, 391, "Hình 2.1. Sơ đồ quan hệ thực thể (ERD) — 3 cơ sở dữ liệu theo mô hình Database-per-Service")

    # Hình 2.6 Audit sequence (para 528) → 15-sequence-audit-pipeline.png
    replace_image_in_paragraph(doc, 528, DIAGRAMS_DIR / "15-sequence-audit-pipeline.png")
    update_paragraph_text(doc, 529, "Hình 2.6. Biểu đồ tuần tự Audit Pipeline (Hybrid Orchestration + Choreography)")

    # Hình 2.8 Architecture (para 565) → 16-component-grpc-bullmq.png
    replace_image_in_paragraph(doc, 565, DIAGRAMS_DIR / "16-component-grpc-bullmq.png")
    update_paragraph_text(doc, 566, "Hình 2.8. Sơ đồ giao tiếp giữa các microservice (gRPC + BullMQ + Redis Pub/Sub)")

    # ============================================================
    # STEP 2: Update descriptions for replaced diagrams
    # ============================================================
    print("\n--- Step 2: Update descriptions ---")

    # Update ERD description (para 389 is the text before ERD image)
    # Find and update the description text around para 389
    for i in range(388, 395):
        p = doc.paragraphs[i]
        if 'Cơ sở dữ liệu' in p.text or 'database' in p.text.lower() or 'PostgreSQL' in p.text:
            break

    # Update audit sequence description (para 527 area)
    for i in range(525, 535):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if 'Khi người dùng nhấn nút' in text or 'luồng chính' in text:
            break

    # Update architecture description (para 564)
    p564 = doc.paragraphs[564]
    if 'REST API' in p564.text or 'Microservices' in p564.text:
        for run in p564.runs:
            run.text = ""
        p564.runs[0].text = (
            "Hệ thống được thiết kế theo mô hình Microservices với 5 service độc lập, "
            "giao tiếp với nhau qua gRPC (đồng bộ) và BullMQ + Redis Pub/Sub (bất đồng bộ). "
            "Mỗi service sở hữu database riêng theo nguyên tắc Database-per-Service."
        )
        print(f"  Updated architecture description at para 564")

    # ============================================================
    # STEP 3: Add 3 new diagram sections after architecture
    # ============================================================
    print("\n--- Step 3: Add new diagram sections ---")

    # Find the right insertion point: after Hình 2.9 (DFD) section, before 2.5.2
    # We'll insert after para 587 (Hình 2.9 caption) and before para 588 (2.5.2)
    # Actually, let's insert after the architecture section (after para 566 Hình 2.8 caption)
    # and before 6.1 Kiến Trúc Tổng Thể

    # Better: Insert new figures after the DFD section (para 587), before 2.5.2
    # This puts them at the end of the architecture overview section

    insert_after = doc.paragraphs[587]._element  # After Hình 2.9 caption

    # --- Hình 2.12: Rule Engine Class Diagram (19) ---
    # Add in reverse order since each addnext goes right after the anchor

    # Empty line
    spacer3 = add_text_paragraph(doc, insert_after, "", style='Normal')

    # Description
    desc3 = add_text_paragraph(
        doc, spacer3,
        "Sơ đồ lớp mô tả kiến trúc SEO Rule Engine theo Strategy Pattern. "
        "Interface ISeoRule định nghĩa phương thức check() cho từng rule. "
        "RuleRegistry quản lý đăng ký/tra cứu rule, RuleRunner thực thi toàn bộ rule trên PageData, "
        "và ScoreCalculator tính điểm trung bình có trọng số rồi phân loại (excellent/good/fair/poor). "
        "Hệ thống bao gồm 20 rule thuộc 6 category: META (4), HEADINGS (2), IMAGES (2), LINKS (2), "
        "PERFORMANCE (1) và TECHNICAL (9).",
        style='Normal'
    )

    # Caption
    cap3 = add_caption_paragraph(doc, desc3, "Hình 2.12. Sơ đồ lớp SEO Rule Engine (Strategy Pattern — 20 rules)")

    # Image
    img3 = add_image_paragraph(doc, cap3, DIAGRAMS_DIR / "19-class-rule-engine.png", width_inches=6.2)

    # Heading
    head3 = add_text_paragraph(doc, img3, "2.5.5. Sơ đồ lớp SEO Rule Engine", style='Heading 3')

    # --- Hình 2.11: Crawler Activity Diagram (18) ---
    spacer2 = add_text_paragraph(doc, head3, "", style='Normal')

    desc2 = add_text_paragraph(
        doc, spacer2,
        "Sơ đồ hoạt động mô tả luồng xử lý của Crawler Service khi nhận yêu cầu crawl URL. "
        "Quá trình bắt đầu bằng kiểm tra SSRF (chặn localhost, IP nội bộ, DNS rebinding), "
        "tiếp theo kiểm tra cache Redis. Nếu cache miss, hệ thống sử dụng Cheerio cho trang tĩnh "
        "(~200ms) và tự động fallback sang Playwright cho SPA (~3-10s). "
        "Cuối cùng, Lighthouse được chạy để thu thập Core Web Vitals (LCP, INP, CLS) với cache 1 giờ.",
        style='Normal'
    )

    cap2 = add_caption_paragraph(doc, desc2, "Hình 2.11. Sơ đồ hoạt động Crawler Service (SSRF → Cheerio/Playwright → Lighthouse)")

    img2 = add_image_paragraph(doc, cap2, DIAGRAMS_DIR / "18-activity-crawler.png", width_inches=5.5)

    head2 = add_text_paragraph(doc, img2, "2.5.4b. Sơ đồ hoạt động Crawler Service", style='Heading 3')

    # --- Hình 2.10: Docker Deployment (17) ---
    spacer1 = add_text_paragraph(doc, head2, "", style='Normal')

    desc1 = add_text_paragraph(
        doc, spacer1,
        "Sơ đồ triển khai mô tả 9 container Docker Compose của hệ thống: "
        "5 service NestJS (Gateway port 3000/50051, Crawler port 50052, SEO Analyzer port 50053, "
        "Keyword Analyzer port 50054, Report port 3004/50055), "
        "3 PostgreSQL 16 instances (Gateway DB port 5432, Analyzer DB port 5433, Report DB port 5434), "
        "và 1 Redis 7 (port 6379) phục vụ cache, BullMQ job queue và Pub/Sub events. "
        "Toàn bộ hệ thống khởi động bằng lệnh docker-compose up -d.",
        style='Normal'
    )

    cap1 = add_caption_paragraph(doc, desc1, "Hình 2.10. Sơ đồ triển khai Docker Compose (9 containers)")

    img1 = add_image_paragraph(doc, cap1, DIAGRAMS_DIR / "17-deployment-docker.png", width_inches=6.2)

    head1 = add_text_paragraph(doc, img1, "2.5.4a. Sơ đồ triển khai Docker Compose", style='Heading 3')

    # ============================================================
    # STEP 4: Update DANH SÁCH HÌNH
    # ============================================================
    print("\n--- Step 4: Update DANH SÁCH HÌNH ---")

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()

        # Update existing figure descriptions
        if text == "Hình 2.1. Sơ đồ quan hệ thực thể (ERD)":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = "Hình 2.1. Sơ đồ quan hệ thực thể (ERD) — 3 cơ sở dữ liệu (Database-per-Service)"
            print(f"  Updated figure list entry at para {i}")

        elif text == "Hình 2.6. Biểu đồ tuần tự chức năng phân tích SEO":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = "Hình 2.6. Biểu đồ tuần tự Audit Pipeline (Hybrid Orchestration + Choreography)"
            print(f"  Updated figure list entry at para {i}")

        elif text == "Hình 2.8. Sơ đồ kiến trúc microservices tổng thể":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = "Hình 2.8. Sơ đồ giao tiếp microservice (gRPC + BullMQ + Redis Pub/Sub)"
            print(f"  Updated figure list entry at para {i}")

        # Add new figure entries after Hình 2.9
        elif text.startswith("Hình 2.9."):
            # Add 3 new entries after this one
            anchor = p._element
            e3 = add_text_paragraph(doc, anchor,
                "Hình 2.12. Sơ đồ lớp SEO Rule Engine (Strategy Pattern — 20 rules)",
                style='Normal')
            e2 = add_text_paragraph(doc, anchor,
                "Hình 2.11. Sơ đồ hoạt động Crawler Service (SSRF → Cheerio/Playwright → Lighthouse)",
                style='Normal')
            e1 = add_text_paragraph(doc, anchor,
                "Hình 2.10. Sơ đồ triển khai Docker Compose (9 containers)",
                style='Normal')
            print(f"  Added 3 new figure list entries after para {i}")
            break

    # ============================================================
    # STEP 5: Save
    # ============================================================
    print(f"\n--- Saving to {OUTPUT_DOCX} ---")
    doc.save(str(OUTPUT_DOCX))
    print(f"Done! File saved: {OUTPUT_DOCX}")
    print(f"File size: {OUTPUT_DOCX.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()

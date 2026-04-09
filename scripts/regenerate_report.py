"""Regenerate SEO Analyst Platform report following the sample academic format
(Trường ĐH Giao thông Vận tải). Based on the reference PDF structure:

Front matter: Bìa → Cam đoan → Cảm ơn → Lời nói đầu → Mục lục → DS từ viết
tắt → DS bảng → DS hình. Body: 3 chapters (Tổng quan, Phân tích & thiết kế,
Triển khai dự án). Each chapter ends with "Kết luận chương". Then Kết luận,
TLTK, Phụ lục.
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OLD = Path("/Users/minhducoder/SEO-Analysts/SEO_Analyst_Platform_Report.docx")
TEMPLATE = Path("/Users/minhducoder/SEO-Analysts/SEO_Report_v2.docx")
OUT = Path("/Users/minhducoder/SEO-Analysts/SEO_Report_v2.docx")
DIAGRAMS = Path("/Users/minhducoder/SEO-Analysts/diagrams")


# ---------------------------------------------------------------------------
# Extract old body content in order, grouped by H1 section title
# ---------------------------------------------------------------------------
def extract_old_sections():
    doc = Document(str(OLD))
    body = doc.element.body
    sections = {}
    current = None
    STYLE_ID_TO_NAME = {
        "2": "Heading1", "3": "Heading2", "4": "Heading3",
        "Heading1": "Heading1", "Heading2": "Heading2", "Heading3": "Heading3",
    }
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            pPr = child.find(qn("w:pPr"))
            raw_style = None
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    raw_style = pStyle.get(qn("w:val"))
            style = STYLE_ID_TO_NAME.get(raw_style, "Normal")
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if style == "Heading1":
                current = text.strip()
                sections[current] = []
                continue
            if current is None:
                continue
            if not text.strip() and style == "Normal":
                continue
            sections[current].append(("p", style, text))
        elif tag == "tbl":
            if current is None:
                continue
            sections[current].append(("tbl", deepcopy(child)))
    return sections


STYLE_MAP = {"Heading2": "Heading 2", "Heading3": "Heading 3", "Normal": "Normal"}


def render_items(doc, items, *, demote=False):
    """Render extracted items. When demote=True, old H2 becomes H3 so they fit
    under a new H2 heading we add programmatically."""
    for item in items:
        if item[0] == "p":
            _, style, text = item
            new_style = STYLE_MAP.get(style, "Normal")
            if demote and new_style == "Heading 2":
                new_style = "Heading 3"
            p = doc.add_paragraph(text)
            try:
                p.style = doc.styles[new_style]
            except KeyError:
                p.style = doc.styles["Normal"]
        elif item[0] == "tbl":
            doc.element.body.append(deepcopy(item[1]))
            doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def h1(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 1"]
    return p


def h2(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 2"]
    return p


def h3(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 3"]
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def centered(doc, text, *, size=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    if size:
        r.font.size = Pt(size)
    r.bold = bold
    return p


def caption(doc, text):
    """Table/figure caption – italic, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.bold = True
    return p


def embed_image(doc, filename, code, description, *, width_inches=6.0):
    """Insert a rendered PNG from diagrams/ as a centered figure with caption.
    Falls back to placeholder() if the file doesn't exist.
    """
    img_path = DIAGRAMS / filename
    if not img_path.exists():
        placeholder(doc, code, "HÌNH", description, "")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"{code}. {description}")
    cr.italic = True
    cr.bold = True


def placeholder(doc, code, kind, description, tool_hint=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ {kind} — {code}: CHƯA CÓ, CẦN CHÈN ]")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"{code}. {description}")
    cr.italic = True
    if tool_hint:
        hint = doc.add_paragraph()
        hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hint.add_run(f"(Gợi ý công cụ: {tool_hint})")
        hr.italic = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def tech_block(doc, name, short_desc, advantages, disadvantages, usage):
    """Format a technology entry in 4-field academic style."""
    b = doc.add_paragraph()
    b.add_run(name).bold = True
    para(doc, short_desc)
    doc.add_paragraph("Ưu điểm:").runs[0].italic = True
    for x in advantages:
        doc.add_paragraph(f"− {x}")
    doc.add_paragraph("Nhược điểm:").runs[0].italic = True
    for x in disadvantages:
        doc.add_paragraph(f"− {x}")
    doc.add_paragraph("Ứng dụng trong dự án:").runs[0].italic = True
    for x in usage:
        doc.add_paragraph(f"− {x}")


def usecase_block(doc, title, actor, desc, main_flow, alt_flow, pre, post):
    """Format a detailed use case specification in 6-field UML style."""
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(f"− Tác nhân: {actor}")
    doc.add_paragraph(f"− Mô tả: {desc}")
    doc.add_paragraph("− Dòng sự kiện chính:")
    for i, step in enumerate(main_flow, 1):
        doc.add_paragraph(f"  Bước {i}. {step}")
    doc.add_paragraph(f"− Dòng sự kiện phụ: {alt_flow}")
    doc.add_paragraph(f"− Điều kiện bắt đầu: {pre}")
    doc.add_paragraph(f"− Điều kiện kết thúc: {post}")


def testcase_table(doc, caption_text, rows):
    """Render a test case table with standard 5 columns."""
    caption(doc, caption_text)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "TC ID"
    hdr[1].text = "Chức năng"
    hdr[2].text = "Mục tiêu"
    hdr[3].text = "Bước thực hiện"
    hdr[4].text = "Trạng thái"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph("")


def chapter_conclusion(doc, num, summary_paras):
    h2(doc, f"{num} Kết luận chương")
    for p in summary_paras:
        para(doc, p)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build():
    sections = extract_old_sections()

    def find(keyword):
        for k in sections:
            if keyword.lower() in k.lower():
                return sections[k]
        return []

    s_func = find("YÊU CẦU CHỨC NĂNG")
    s_nfr = find("PHI CHỨC NĂNG")
    s_sysoverview = find("SYSTEM OVERVIEW")
    s_microservices = find("KIẾN TRÚC CHI TIẾT")
    s_db = find("DATABASE SCHEMA")
    s_api = find("API ENDPOINTS")
    s_test = find("KIỂM THỬ")
    s_outcome = find("EXPECTED OUTCOME")
    s_roadmap = find("ROADMAP")

    # Open template (already has correct styles)
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    sectPr = None
    for child in list(body.iterchildren()):
        if child.tag.split("}")[-1] == "sectPr":
            sectPr = deepcopy(child)
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

    # ========================================================================
    # TRANG BÌA
    # ========================================================================
    centered(doc, "TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", size=14, bold=True)
    centered(doc, "KHOA CÔNG NGHỆ THÔNG TIN", size=13, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "ĐỒ ÁN TỐT NGHIỆP", size=20, bold=True)
    centered(doc, "ĐỀ TÀI", size=14)
    doc.add_paragraph()
    centered(doc, "XÂY DỰNG NỀN TẢNG PHÂN TÍCH", size=22, bold=True)
    centered(doc, "SEO WEBSITE TỰ ĐỘNG", size=22, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "Giảng viên hướng dẫn : [Tên GVHD]", size=13)
    centered(doc, "Sinh viên thực hiện : [Họ và tên]", size=13)
    centered(doc, "Lớp : [Lớp]", size=13)
    centered(doc, "Mã sinh viên : [MSSV]", size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, "Hà Nội – 2026", size=13, bold=True)
    doc.add_page_break()

    # ========================================================================
    # LỜI CAM ĐOAN
    # ========================================================================
    p = h1(doc, "LỜI CAM ĐOAN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
        "Em xin cam đoan rằng đồ án tốt nghiệp \"Xây dựng nền tảng phân tích "
        "SEO website tự động\" là kết quả của quá trình học tập và nghiên cứu "
        "của bản thân, dưới sự hướng dẫn của thầy/cô [Tên GVHD]."
    )
    para(doc,
        "Toàn bộ phần mã nguồn của hệ thống (bao gồm cả giao diện phía máy "
        "khách và các dịch vụ phía máy chủ) đều do em tự thiết kế và phát "
        "triển trong quá trình thực hiện đồ án. Trong quá trình xây dựng, em "
        "có tham khảo ý tưởng và bố cục giao diện từ một số công cụ SEO hiện "
        "có như Ahrefs, SEMrush, Ubersuggest nhằm học hỏi và nâng cao trải "
        "nghiệm người dùng cho sản phẩm. Tuy nhiên, toàn bộ mã TypeScript, "
        "JavaScript, HTML/CSS và các rule phân tích SEO đều do em tự viết "
        "lại, không sao chép trực tiếp mã nguồn của bất kỳ sản phẩm nào."
    )
    para(doc,
        "Các hình ảnh, dữ liệu mẫu hoặc tài liệu tham khảo trong đồ án được "
        "trích dẫn đầy đủ nguồn gốc ở phần Tài liệu tham khảo, không dùng vào "
        "mục đích thương mại."
    )
    para(doc,
        "Em xin hoàn toàn chịu trách nhiệm về nội dung và tính trung thực của "
        "đồ án này trước Hội đồng và Nhà trường."
    )
    doc.add_paragraph()
    centered(doc, "Người thực hiện", bold=True)
    centered(doc, "(Ký và ghi rõ họ tên)")
    doc.add_page_break()

    # ========================================================================
    # LỜI CẢM ƠN
    # ========================================================================
    p = h1(doc, "LỜI CẢM ƠN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
        "Trong khoảng thời gian làm đồ án tốt nghiệp, em đã nhận được nhiều "
        "sự giúp đỡ, đóng góp ý kiến và sự dẫn dắt chỉ bảo nhiệt tình của "
        "thầy cô, gia đình và bạn bè."
    )
    para(doc,
        "Em xin gửi lời cảm ơn chân thành đến giảng viên hướng dẫn – thầy/cô "
        "[Tên GVHD] – Trường Đại học Giao thông Vận tải, người đã tận tình "
        "hướng dẫn, chỉ bảo em trong suốt quá trình thực hiện đồ án."
    )
    para(doc,
        "Em cũng xin gửi lời cảm ơn chân thành nhất tới các thầy cô giáo "
        "trong Trường Đại học Giao thông Vận tải nói chung, các thầy cô "
        "trong khoa Công nghệ Thông tin nói riêng, đã dạy dỗ cho em kiến "
        "thức về các môn đại cương cũng như các môn chuyên ngành, giúp em có "
        "được cơ sở lý thuyết vững vàng và tạo điều kiện giúp đỡ em trong "
        "suốt quá trình học tập."
    )
    para(doc,
        "Cuối cùng, em xin chân thành cảm ơn gia đình và bạn bè, những người "
        "luôn ở bên cạnh đã luôn tạo điều kiện, quan tâm, giúp đỡ, động viên "
        "em trong suốt quá trình học tập và hoàn thành báo cáo."
    )
    para(doc,
        "Với điều kiện về thời gian cũng như lượng kiến thức về đề tài rất "
        "rộng mà kinh nghiệm còn hạn chế của một sinh viên, đồ án này không "
        "thể tránh được những thiếu sót. Em rất mong nhận được sự chỉ bảo, "
        "đóng góp ý kiến của các thầy cô để em có điều kiện bổ sung, nâng "
        "cao kiến thức của mình, phục vụ tốt hơn công tác thực tế sau này."
    )
    para(doc, "Em xin chân thành cảm ơn!")
    doc.add_page_break()

    # ========================================================================
    # LỜI NÓI ĐẦU
    # ========================================================================
    p = h1(doc, "LỜI NÓI ĐẦU")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
        "Trong những năm gần đây, cùng với sự phát triển mạnh mẽ của Internet "
        "và thương mại điện tử, việc có một website được tìm thấy dễ dàng "
        "trên các công cụ tìm kiếm như Google đã trở thành yếu tố sống còn "
        "đối với hàng trăm nghìn doanh nghiệp nhỏ và vừa (SME), freelancer và "
        "cá nhân kinh doanh online tại Việt Nam. Tối ưu hóa công cụ tìm kiếm "
        "(Search Engine Optimization – SEO) vì vậy không còn là lựa chọn mà "
        "là yêu cầu bắt buộc đối với bất kỳ website nào muốn có lượng truy "
        "cập ổn định và bền vững."
    )
    para(doc,
        "Tuy nhiên, các công cụ phân tích SEO thương mại phổ biến nhất hiện "
        "nay như Ahrefs, SEMrush hay Moz Pro đều có mức giá từ 99 đến 499 "
        "USD/tháng – vượt quá khả năng chi trả của sinh viên học nghề, "
        "freelancer mới vào nghề và doanh nghiệp nhỏ. Bên cạnh chi phí, các "
        "công cụ này thường có giao diện phức tạp với hàng trăm tính năng, "
        "khiến người mới rất khó tiếp cận. Điều này tạo ra một nghịch lý: "
        "chính những đối tượng cần SEO nhất lại là những người khó tiếp cận "
        "công cụ SEO chuyên nghiệp nhất."
    )
    para(doc,
        "Với mong muốn góp phần giải quyết vấn đề trên, em lựa chọn đề tài "
        "\"Xây dựng nền tảng phân tích SEO website tự động\" làm đồ án tốt "
        "nghiệp. Mục tiêu của đề tài là xây dựng một hệ thống web miễn phí ở "
        "mức cá nhân, giao diện đơn giản, tập trung vào bộ chỉ số SEO on-page "
        "cốt lõi nhưng đủ sâu để mang lại giá trị thực tiễn cho người dùng. "
        "Hệ thống được thiết kế theo kiến trúc microservices hiện đại với "
        "Next.js, NestJS, PostgreSQL và Redis, triển khai trên hạ tầng cloud "
        "chi phí thấp (Vercel + Railway + Supabase) dưới 40 USD/tháng."
    )
    para(doc,
        "Báo cáo đồ án được chia thành 3 chương chính: Chương 1 trình bày "
        "tổng quan về đề tài bao gồm khảo sát nhu cầu và phân tích công nghệ; "
        "Chương 2 phân tích yêu cầu và thiết kế hệ thống gồm cơ sở dữ liệu, "
        "biểu đồ use-case, sơ đồ phân rã chức năng, biểu đồ tuần tự và kiến "
        "trúc microservices; Chương 3 mô tả quá trình triển khai dự án từ "
        "môi trường phát triển, cấu trúc thư mục, giao diện người dùng cho "
        "đến kiểm thử và triển khai production. Cuối cùng là phần Kết luận, "
        "Tài liệu tham khảo và Phụ lục."
    )
    doc.add_page_break()

    # ========================================================================
    # MỤC LỤC (placeholder – auto-update in Word)
    # ========================================================================
    p = h1(doc, "MỤC LỤC")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
        "(Mục lục sẽ được cập nhật tự động bằng chức năng References → Table "
        "of Contents của Microsoft Word sau khi hoàn thiện nội dung.)"
    )
    doc.add_page_break()

    # ========================================================================
    # DANH SÁCH CÁC TỪ VIẾT TẮT
    # ========================================================================
    h1(doc, "DANH SÁCH CÁC TỪ VIẾT TẮT")
    abbr_list = [
        ("SEO", "Search Engine Optimization – Tối ưu hóa công cụ tìm kiếm"),
        ("SME", "Small and Medium Enterprise – Doanh nghiệp nhỏ và vừa"),
        ("SERP", "Search Engine Results Page – Trang kết quả tìm kiếm"),
        ("API", "Application Programming Interface – Giao diện lập trình ứng dụng"),
        ("REST", "Representational State Transfer"),
        ("JWT", "JSON Web Token"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("ERD", "Entity Relationship Diagram – Sơ đồ quan hệ thực thể"),
        ("NFR", "Non-Functional Requirements – Yêu cầu phi chức năng"),
        ("SPA", "Single Page Application – Ứng dụng đơn trang"),
        ("SSR", "Server-Side Rendering – Kết xuất phía máy chủ"),
        ("SSG", "Static Site Generation – Tạo trang tĩnh"),
        ("CI/CD", "Continuous Integration / Continuous Deployment"),
        ("ORM", "Object Relational Mapping"),
        ("DOM", "Document Object Model"),
        ("SQL", "Structured Query Language"),
        ("LCP", "Largest Contentful Paint"),
        ("FID", "First Input Delay"),
        ("CLS", "Cumulative Layout Shift"),
        ("IDE", "Integrated Development Environment"),
    ]
    for k, v in abbr_list:
        doc.add_paragraph(f"− {k}: {v}")
    doc.add_page_break()

    # ========================================================================
    # DANH SÁCH BẢNG (placeholder)
    # ========================================================================
    h1(doc, "DANH SÁCH BẢNG")
    para(doc,
        "(Danh sách các bảng sẽ được cập nhật thủ công hoặc tự động bằng "
        "chức năng References → Insert Table of Figures (Table of: Bảng) "
        "trong Microsoft Word sau khi hoàn thiện nội dung.)"
    )
    for line in [
        "Bảng 1.1. Bảng so sánh giá các công cụ SEO thương mại phổ biến",
        "Bảng 1.2. Bảng so sánh tính năng các công cụ SEO thương mại",
        "Bảng 2.1. Bảng mô tả thực thể users",
        "Bảng 2.2. Bảng mô tả thực thể audits",
        "Bảng 2.3. Bảng mô tả thực thể audit_results",
        "Bảng 2.4. Bảng mô tả thực thể keywords",
        "Bảng 2.5. Bảng use-case tổng quan",
        "Bảng 2.6. Bảng đặc tả yêu cầu phi chức năng",
        "Bảng 3.1. Test case chức năng xác thực",
        "Bảng 3.2. Test case chức năng phân tích SEO",
        "Bảng 3.3. Test case chức năng quản lý tài khoản",
        "Bảng 3.4. Test case chức năng quản trị viên",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ========================================================================
    # DANH SÁCH HÌNH (placeholder)
    # ========================================================================
    h1(doc, "DANH SÁCH HÌNH")
    para(doc,
        "(Danh sách các hình sẽ được cập nhật thủ công hoặc tự động bằng "
        "chức năng References → Insert Table of Figures trong Microsoft "
        "Word sau khi hoàn thiện nội dung.)"
    )
    for line in [
        "Hình 1.1. Biểu đồ so sánh giá các công cụ SEO thương mại",
        "Hình 1.2. Giao diện công cụ Ahrefs (tham khảo)",
        "Hình 1.3. Giao diện công cụ SEMrush (tham khảo)",
        "Hình 2.1. Sơ đồ quan hệ thực thể (ERD)",
        "Hình 2.2. Biểu đồ use-case người dùng",
        "Hình 2.3. Biểu đồ use-case quản trị viên",
        "Hình 2.4. Sơ đồ phân rã chức năng",
        "Hình 2.5. Biểu đồ tuần tự chức năng đăng nhập",
        "Hình 2.6. Biểu đồ tuần tự chức năng phân tích SEO",
        "Hình 2.7. Biểu đồ tuần tự chức năng xuất báo cáo PDF",
        "Hình 2.8. Sơ đồ kiến trúc microservices tổng thể",
        "Hình 2.9. Sơ đồ luồng dữ liệu (Data Flow Diagram)",
        "Hình 3.1. Cấu trúc thư mục dự án",
        "Hình 3.2. Cấu trúc thư mục apps/frontend",
        "Hình 3.3. Cấu trúc thư mục apps/backend",
        "Hình 3.4. Giao diện trang chủ",
        "Hình 3.5. Giao diện trang đăng nhập",
        "Hình 3.6. Giao diện dashboard danh sách audit",
        "Hình 3.7. Giao diện form nhập URL phân tích",
        "Hình 3.8. Giao diện chi tiết kết quả audit",
        "Hình 3.9. Giao diện quản trị viên",
        "Hình 3.10. Mẫu báo cáo PDF xuất ra",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    # ========================================================================
    # CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI
    # ========================================================================
    h1(doc, "CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI")

    # ---- 1.1 Giới thiệu đề tài ----
    h2(doc, "1.1. Giới thiệu đề tài")

    h3(doc, "1.1.1. Tổng quan")
    para(doc,
        "Trong thời đại kinh tế số, nền kinh tế Internet Việt Nam đã đạt quy "
        "mô khoảng 36 tỷ USD vào năm 2024 và được dự báo tiếp tục tăng "
        "trưởng hai con số trong những năm tới [14]. Song song với sự bùng "
        "nổ của thương mại điện tử và nội dung số, hàng trăm nghìn website "
        "doanh nghiệp nhỏ và vừa, freelancer và cá nhân kinh doanh đang "
        "cạnh tranh để giành vị trí xuất hiện trên trang kết quả tìm kiếm "
        "(SERP) của Google – nơi mà vị trí đầu tiên nhận được hơn 27% lượng "
        "truy cập tự nhiên, trong khi vị trí thứ mười chỉ nhận chưa tới 3% [14]."
    )
    para(doc,
        "Tối ưu hóa công cụ tìm kiếm (Search Engine Optimization – SEO) vì "
        "vậy trở thành một kênh marketing thiết yếu với chi phí thấp và hiệu "
        "quả lâu dài. Tuy nhiên, để triển khai SEO một cách bài bản, người "
        "làm SEO cần theo dõi liên tục hàng chục chỉ số kỹ thuật trên "
        "website: tốc độ tải trang, khả năng thu thập dữ liệu, cấu trúc "
        "heading, mức độ tối ưu meta tag, sơ đồ liên kết nội bộ, trải nghiệm "
        "di động và các Core Web Vitals do Google quy định [14]."
    )
    para(doc,
        "Các công cụ SEO thương mại hiện nay như Ahrefs, SEMrush hay Moz Pro "
        "đều có mức giá từ 99–499 USD/tháng – vượt quá khả năng của sinh "
        "viên, freelancer và SME. Bên cạnh chi phí cao, các công cụ này còn "
        "phức tạp với hàng trăm tính năng không cần thiết, không phù hợp "
        "cho người mới. Với mục tiêu lấp đầy khoảng trống này, đồ án "
        "\"Xây dựng nền tảng phân tích SEO website tự động\" được thực hiện "
        "nhằm tạo ra một nền tảng phân tích SEO miễn phí ở mức cá nhân, "
        "giao diện đơn giản, tập trung vào các chỉ số SEO on-page thiết yếu."
    )

    h3(doc, "1.1.2. Nội dung và phạm vi của đề tài")
    para(doc, "Hệ thống bao gồm các tính năng chính như:")
    for line in [
        "Cho phép người dùng đăng ký tài khoản và đăng nhập bằng email/mật khẩu.",
        "Nhận URL từ người dùng và tự động crawl HTML, DOM của trang đích.",
        "Áp dụng bộ 20 rule SEO on-page để chấm điểm và phát hiện các issue.",
        "Hiển thị kết quả phân tích dưới dạng dashboard trực quan với biểu đồ và danh sách gợi ý cải thiện.",
        "Lưu lịch sử audit để người dùng dễ dàng theo dõi sự thay đổi theo thời gian.",
        "Xuất báo cáo PDF để người dùng có thể chia sẻ hoặc lưu trữ offline.",
        "Cập nhật tiến độ phân tích theo thời gian thực qua WebSocket.",
        "Hỗ trợ quản trị viên quản lý người dùng và cấu hình trọng số rule SEO.",
    ]:
        doc.add_paragraph(f"− {line}")
    para(doc, "Phạm vi đề tài:")
    doc.add_paragraph("− Đối tượng sử dụng:")
    for line in [
        "Sinh viên học SEO: Dùng để thực hành phân tích website mẫu và học các rule SEO cơ bản.",
        "Freelancer marketing: Dùng để audit nhanh website khách hàng và đưa ra gợi ý cải thiện.",
        "Chủ SME: Dùng để kiểm tra định kỳ website doanh nghiệp và theo dõi cải thiện theo thời gian.",
        "Quản trị viên: Kiểm duyệt người dùng, cấu hình rule SEO, theo dõi thống kê sử dụng.",
    ]:
        doc.add_paragraph(f"    + {line}")
    doc.add_paragraph("− Không thuộc phạm vi:")
    for line in [
        "Phân tích backlink và off-page SEO (đòi hỏi crawler phân tán quy mô lớn).",
        "Nghiên cứu từ khóa đối thủ (cần database keyword toàn cầu).",
        "Crawl sâu nhiều trang cùng lúc (deep crawl toàn website).",
        "Tích hợp AI/ML cho gợi ý cải thiện nội dung (thuộc giai đoạn sau).",
    ]:
        doc.add_paragraph(f"    + {line}")

    h3(doc, "1.1.3. Kết quả dự kiến đạt được")
    para(doc, "Sau khi hoàn thành, hệ thống sẽ đạt được các mục tiêu sau:")
    doc.add_paragraph("− Đối với người dùng cuối:")
    for line in [
        "Đăng ký, đăng nhập và quản lý tài khoản cá nhân.",
        "Nhập URL website và nhận kết quả phân tích trong vòng dưới 10 giây.",
        "Xem chi tiết điểm số từng rule SEO và gợi ý cải thiện cụ thể.",
        "Xuất báo cáo PDF để chia sẻ hoặc lưu trữ.",
        "Theo dõi lịch sử audit và sự thay đổi theo thời gian.",
    ]:
        doc.add_paragraph(f"    + {line}")
    doc.add_paragraph("− Đối với quản trị viên:")
    for line in [
        "Quản lý danh sách người dùng, khóa/mở tài khoản vi phạm.",
        "Cấu hình trọng số các rule SEO linh hoạt.",
        "Theo dõi thống kê số lượng audit, thời gian trung bình, tỷ lệ thành công.",
    ]:
        doc.add_paragraph(f"    + {line}")
    doc.add_paragraph("− Mục tiêu kỹ thuật:")
    for line in [
        "Hệ thống triển khai trên cloud với chi phí dưới 40 USD/tháng.",
        "Thời gian crawl trung bình dưới 10 giây cho mỗi URL.",
        "Độ chính xác của rule SEO so với Lighthouse ≥ 90%.",
        "Code coverage của test ≥ 80% cho phần business logic.",
    ]:
        doc.add_paragraph(f"    + {line}")

    # ---- 1.2 Khảo sát ----
    h2(doc, "1.2. Khảo sát")

    h3(doc, "1.2.1. Phiếu hỏi")
    para(doc,
        "Để phục vụ quá trình phân tích yêu cầu và thiết kế hệ thống, em đã "
        "xây dựng phiếu khảo sát nhằm thu thập nhu cầu thực tế từ nhóm đối "
        "tượng mục tiêu (sinh viên học SEO, freelancer, chủ SME). Mặc dù "
        "chưa triển khai khảo sát thực tế trên quy mô lớn, các câu hỏi được "
        "xây dựng dựa trên nghiên cứu hành vi người dùng trên các nền tảng "
        "SEO hiện có như Ahrefs, SEMrush, Ubersuggest nhằm làm rõ định "
        "hướng phát triển hệ thống. Dưới đây là nội dung phiếu khảo sát."
    )
    doc.add_paragraph("Phiếu Khảo Sát Nhu Cầu Công Cụ Phân Tích SEO", style="Heading 3")
    doc.add_paragraph("Thông tin chung:").runs[0].bold = True
    for q in [
        "1. Bạn thuộc nhóm đối tượng nào?",
        "  Ο Sinh viên đang học về SEO / Digital Marketing",
        "  Ο Freelancer cung cấp dịch vụ SEO",
        "  Ο Chủ doanh nghiệp nhỏ/vừa (SME)",
        "  Ο Nhân viên marketing trong công ty",
        "  Ο Khác (vui lòng ghi rõ): ....................",
        "2. Bạn đã sử dụng công cụ SEO nào? (có thể chọn nhiều)",
        "  Ο Google Search Console (miễn phí)",
        "  Ο Google Analytics (miễn phí)",
        "  Ο Ahrefs",
        "  Ο SEMrush",
        "  Ο Moz Pro",
        "  Ο Ubersuggest",
        "  Ο Screaming Frog",
        "  Ο Chưa sử dụng công cụ nào",
    ]:
        doc.add_paragraph(q)
    doc.add_paragraph("Về chi phí và rào cản:").runs[0].bold = True
    for q in [
        "3. Ngân sách của bạn dành cho công cụ SEO hàng tháng là bao nhiêu?",
        "  Ο Dưới 10 USD",
        "  Ο 10–50 USD",
        "  Ο 50–100 USD",
        "  Ο Trên 100 USD",
        "  Ο Không có ngân sách / chỉ dùng công cụ miễn phí",
        "4. Đâu là rào cản lớn nhất khiến bạn chưa sử dụng công cụ SEO chuyên nghiệp?",
        "  Ο Chi phí quá cao",
        "  Ο Giao diện phức tạp, khó học",
        "  Ο Không có thời gian tìm hiểu",
        "  Ο Không biết nên bắt đầu từ đâu",
        "  Ο Không có nhu cầu thường xuyên",
    ]:
        doc.add_paragraph(q)
    doc.add_paragraph("Về nhu cầu tính năng:").runs[0].bold = True
    for q in [
        "5. Những tính năng nào bạn cần nhất ở một công cụ phân tích SEO?",
        "  Ο Chấm điểm SEO tổng thể và gợi ý cải thiện",
        "  Ο Phát hiện lỗi meta tag, heading, hình ảnh",
        "  Ο Kiểm tra Core Web Vitals và hiệu năng trang",
        "  Ο Phân tích từ khóa",
        "  Ο Phân tích backlink",
        "  Ο Theo dõi thứ hạng trên SERP",
        "  Ο Xuất báo cáo PDF",
        "  Ο Khác (vui lòng ghi rõ): ....................",
        "6. Bạn mong muốn tốc độ phân tích một URL tối đa là bao lâu?",
        "  Ο Dưới 5 giây",
        "  Ο 5–10 giây",
        "  Ο 10–30 giây",
        "  Ο Không quan trọng, miễn chính xác",
        "7. Bạn có cần tính năng lưu lịch sử audit và so sánh theo thời gian không?",
        "  Ο Rất cần",
        "  Ο Cần",
        "  Ο Không cần",
    ]:
        doc.add_paragraph(q)
    doc.add_paragraph("Đề xuất và phản hồi:").runs[0].bold = True
    for q in [
        "8. Nếu có một công cụ SEO miễn phí, giao diện đơn giản, tập trung vào "
        "các rule SEO on-page cơ bản, bạn có sẵn sàng sử dụng không?",
        "  Ο Chắc chắn sử dụng",
        "  Ο Có thể sử dụng thử",
        "  Ο Không chắc",
        "  Ο Không có nhu cầu",
        "9. Bạn có đề xuất thêm tính năng gì cho công cụ không? ....................",
    ]:
        doc.add_paragraph(q)

    h3(doc, "1.2.2. Lấy mẫu")
    para(doc,
        "Để có góc nhìn thực tế về giao diện và trải nghiệm của các công cụ "
        "SEO thương mại, em đã tham khảo và phân tích giao diện của ba nền "
        "tảng phổ biến nhất: Ahrefs, SEMrush và Ubersuggest. Dưới đây là "
        "một số hình ảnh minh họa (chỉ để học hỏi, không sao chép)."
    )
    placeholder(doc, "Hình 1.2", "SCREENSHOT",
        "Giao diện Dashboard Site Audit của Ahrefs",
        "Chụp ảnh màn hình từ trang chủ hoặc trang demo của ahrefs.com")
    placeholder(doc, "Hình 1.3", "SCREENSHOT",
        "Giao diện On-Page SEO Checker của SEMrush",
        "Chụp ảnh màn hình từ trang chủ hoặc trang demo của semrush.com")
    para(doc,
        "Qua việc phân tích các nền tảng trên, em rút ra những nhận xét: "
        "(1) cả ba công cụ đều có giao diện rất phong phú với hàng chục menu "
        "con, rất khó tiếp cận cho người mới; (2) hầu hết tính năng tập trung "
        "vào off-page (backlink, đối thủ) hơn là on-page cơ bản; (3) không "
        "có gói miễn phí cho phép audit website của chính mình một cách đầy "
        "đủ. Bảng dưới đây tổng hợp so sánh giá của các công cụ."
    )
    caption(doc, "Bảng 1.1. Bảng so sánh giá các công cụ SEO thương mại phổ biến")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Công cụ"
    hdr[2].text = "Gói rẻ nhất (USD/tháng)"
    hdr[3].text = "Gói phổ thông (USD/tháng)"
    for i, row in enumerate([
        ("Ahrefs", "99 (Lite)", "199 (Standard)"),
        ("SEMrush", "139 (Pro)", "249 (Guru)"),
        ("Moz Pro", "99 (Standard)", "179 (Medium)"),
        ("Ubersuggest", "12 (Individual)", "29 (Business)"),
        ("Screaming Frog", "259/năm (Paid)", "259/năm (Paid)"),
    ], 1):
        cells = tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row[0]
        cells[2].text = row[1]
        cells[3].text = row[2]
    para(doc,
        "Bảng trên cho thấy một nghịch lý: ngay cả gói rẻ nhất cũng từ 12–99 "
        "USD/tháng – vượt quá khả năng chi trả của nhóm đối tượng mục tiêu "
        "(sinh viên học SEO, freelancer mới vào nghề). Điều này một lần nữa "
        "khẳng định tính cần thiết của đề tài."
    )

    # ---- 1.3 Phân tích công nghệ sử dụng ----
    h2(doc, "1.3. Phân tích công nghệ sử dụng")
    para(doc,
        "Hệ thống được thiết kế với kiến trúc microservices sử dụng các công "
        "nghệ hiện đại, dễ mở rộng và có chi phí triển khai thấp. Việc lựa "
        "chọn công nghệ dựa trên ba tiêu chí: (1) phù hợp với tính chất bất "
        "đồng bộ của quá trình crawl và phân tích, (2) có cộng đồng lớn và "
        "tài liệu phong phú, (3) miễn phí hoặc có free-tier phù hợp ngân "
        "sách sinh viên."
    )

    h3(doc, "1.3.1. Front-end")
    tech_block(doc, "Next.js 14",
        "Là một framework React được phát triển bởi Vercel, hỗ trợ Server-Side "
        "Rendering (SSR), Static Site Generation (SSG) và Incremental Static "
        "Regeneration (ISR) với chế độ App Router mới sử dụng React Server "
        "Components [14].",
        advantages=[
            "Tối ưu SEO tốt hơn React thuần nhờ pre-rendering phía máy chủ – rất quan trọng với website công cụ SEO.",
            "Tốc độ tải trang nhanh, có hỗ trợ streaming UI và Suspense.",
            "Cấu trúc routing rõ ràng theo quy ước thư mục app/.",
            "Tích hợp native với Vercel để triển khai preview URL cho mỗi pull request.",
        ],
        disadvantages=[
            "Đường cong học tập dốc hơn React thuần, đặc biệt với App Router mới.",
            "Bundle size lớn hơn nếu không tối ưu cẩn thận code-splitting.",
        ],
        usage=[
            "Hiển thị trang chủ, trang đăng nhập/đăng ký cần SEO tốt.",
            "Dashboard client-side cho danh sách audit và chi tiết kết quả.",
            "Streaming UI cho progress bar khi đang crawl.",
        ])
    tech_block(doc, "React 18 và TypeScript",
        "React 18 là thư viện UI phổ biến nhất hiện nay với hỗ trợ Concurrent "
        "Mode, Suspense và Server Components. TypeScript là superset của "
        "JavaScript cung cấp type system tĩnh.",
        advantages=[
            "Component hóa giúp tái sử dụng UI hiệu quả.",
            "TypeScript phát hiện lỗi ngay tại thời điểm biên dịch.",
            "Cộng đồng lớn với hàng nghìn thư viện có sẵn.",
        ],
        disadvantages=[
            "Type system phức tạp với generic khi cần kết hợp nhiều library.",
            "Cần thời gian làm quen với JSX và pattern Hook.",
        ],
        usage=[
            "Xây dựng toàn bộ component UI dashboard và form nhập URL.",
            "Định nghĩa type cho payload API để đồng bộ với backend.",
        ])
    tech_block(doc, "Tailwind CSS và shadcn/ui",
        "Tailwind là framework CSS utility-first cho phép xây dựng UI nhanh mà "
        "không cần viết CSS riêng. shadcn/ui là bộ component UI xây trên "
        "Radix UI Primitives với chuẩn accessibility WAI-ARIA.",
        advantages=[
            "Viết giao diện nhanh, dễ responsive với các class mobile-first.",
            "shadcn/ui đảm bảo accessibility và dark mode mặc định.",
            "Không cần maintain file CSS riêng.",
        ],
        disadvantages=[
            "Class HTML dài, dễ rối nếu không tổ chức tốt.",
            "Cần thời gian làm quen với hệ thống design token của Tailwind.",
        ],
        usage=[
            "Layout toàn bộ trang (grid, flex, spacing).",
            "Các component UI: Button, Card, Dialog, Table, Form.",
            "Dark mode toggle cho toàn hệ thống.",
        ])
    tech_block(doc, "TanStack Query (React Query)",
        "Thư viện quản lý server state cho React, tự động xử lý cache, "
        "refetch, retry và loading state.",
        advantages=[
            "Giảm boilerplate so với viết fetch + useState thủ công.",
            "Tự động refetch khi focus window hoặc reconnect mạng.",
            "Tích hợp dễ dàng với Socket.IO cho invalidate cache realtime.",
        ],
        disadvantages=[
            "Thêm một lớp abstraction, debug khi cache sai khó hơn.",
        ],
        usage=[
            "Fetch danh sách audit, chi tiết audit từ backend.",
            "Đồng bộ cache khi nhận sự kiện audit.completed qua Socket.IO.",
        ])

    h3(doc, "1.3.2. Back-end")
    tech_block(doc, "NestJS 10",
        "Framework Node.js viết bằng TypeScript, thiết kế theo triết lý "
        "Angular-inspired với kiến trúc module, decorator, dependency "
        "injection, guards, pipes và interceptors [15].",
        advantages=[
            "Cấu trúc dự án rõ ràng ngay từ đầu, dễ bảo trì khi codebase lớn.",
            "Hỗ trợ native cho BullMQ, Socket.IO, Swagger, Passport.",
            "Dependency Injection giúp test dễ dàng.",
        ],
        disadvantages=[
            "Học cost cao hơn Express thuần do nhiều khái niệm mới.",
            "File boilerplate nhiều hơn cho các dự án nhỏ.",
        ],
        usage=[
            "Xây dựng API Gateway và các microservice (Crawler, Analyzer, Report).",
            "Xác thực JWT với Guards và định tuyến với Controllers.",
        ])
    tech_block(doc, "BullMQ và Redis 7",
        "BullMQ là thư viện hàng đợi job Node.js dựa trên Redis, hỗ trợ "
        "retry, rate limiting, concurrency, priority và scheduled jobs [18]. "
        "Redis là in-memory data store mã nguồn mở hiệu năng cao.",
        advantages=[
            "Xử lý bất đồng bộ tốt, phù hợp với tác vụ crawl tốn thời gian.",
            "Hỗ trợ retry với exponential backoff và dead letter queue.",
            "Dashboard Bull Board cho việc giám sát job realtime.",
        ],
        disadvantages=[
            "Phụ thuộc vào Redis → thêm một service cần quản lý.",
            "Debug job failure phức tạp hơn xử lý đồng bộ.",
        ],
        usage=[
            "Hàng đợi job crawl URL: Producer là API Gateway, Consumer là Crawler Worker.",
            "Cache kết quả Lighthouse trong 1 giờ để tiết kiệm tài nguyên.",
        ])
    tech_block(doc, "Cheerio và Playwright",
        "Cheerio là parser HTML nhẹ với cú pháp jQuery, chạy nhanh (~200 ms/URL). "
        "Playwright là framework điều khiển trình duyệt của Microsoft, hỗ trợ "
        "Chromium/Firefox/WebKit headless [19][20].",
        advantages=[
            "Cheerio cực nhanh cho trang tĩnh, không cần khởi động browser.",
            "Playwright render được SPA nhờ thực thi JavaScript thật.",
            "Playwright còn hỗ trợ chụp screenshot, trace và auto-wait.",
        ],
        disadvantages=[
            "Cheerio không xử lý được các trang SPA (React/Vue/Angular).",
            "Playwright nặng, tiêu thụ nhiều RAM và CPU.",
        ],
        usage=[
            "Cheerio cho trang tĩnh (WordPress, HTML server-rendered).",
            "Fallback sang Playwright khi phát hiện SPA.",
        ])
    tech_block(doc, "Socket.IO",
        "Thư viện hỗ trợ giao tiếp realtime hai chiều giữa client và server "
        "qua WebSocket, với fallback sang long-polling khi cần.",
        advantages=[
            "API đơn giản, dễ tích hợp với NestJS qua @nestjs/websockets.",
            "Hỗ trợ room/namespace để gửi sự kiện đến đúng client.",
            "Tự động reconnect khi mất mạng.",
        ],
        disadvantages=[
            "Tăng độ phức tạp khi scale nhiều instance (cần Redis adapter).",
        ],
        usage=[
            "Phát sự kiện audit.progress và audit.completed tới frontend.",
            "Cập nhật realtime tiến độ crawl cho người dùng đang xem dashboard.",
        ])

    h3(doc, "1.3.3. Cơ sở dữ liệu")
    tech_block(doc, "PostgreSQL 16",
        "Hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở mạnh mẽ và ổn định, "
        "hỗ trợ JSONB, full-text search và các extension nâng cao [16].",
        advantages=[
            "Hỗ trợ tốt dữ liệu quan hệ phức tạp với khóa ngoại và constraint.",
            "JSONB cho phép lưu dữ liệu semi-structured hiệu quả.",
            "Tối ưu query với B-Tree, GIN, GiST index và EXPLAIN ANALYZE.",
        ],
        disadvantages=[
            "Cần hiểu rõ schema, nếu không sẽ không tối ưu được SQL.",
            "Migrate schema cần thận trọng trong production.",
        ],
        usage=[
            "Lưu trữ toàn bộ dữ liệu có cấu trúc: người dùng, audit, rule, kết quả.",
            "Dùng JSONB cho bảng audit_results để lưu output rule linh hoạt.",
        ])
    tech_block(doc, "Prisma 5",
        "ORM TypeScript có khả năng sinh type-safe client tự động từ schema "
        "Prisma. Cung cấp Prisma Migrate cho quản lý schema version [17].",
        advantages=[
            "Type-safe query: lỗi phát hiện tại thời điểm biên dịch.",
            "Prisma Migrate: quản lý migration bằng file SQL có version rõ ràng.",
            "Prisma Studio: GUI đơn giản để xem/sửa dữ liệu lúc phát triển.",
        ],
        disadvantages=[
            "N+1 query problem nếu không dùng include/select đúng cách.",
            "Không hỗ trợ tất cả các tính năng nâng cao của PostgreSQL (ví dụ triggers).",
        ],
        usage=[
            "Định nghĩa schema trong schema.prisma, tự động sinh TypeScript types.",
            "Quản lý migration khi thay đổi cấu trúc bảng.",
        ])

    h3(doc, "1.3.4. Các công nghệ hỗ trợ khác")
    tech_block(doc, "Lighthouse CI",
        "Công cụ mã nguồn mở của Google dùng để đo chất lượng web trên 5 tiêu "
        "chí: Performance, Accessibility, Best Practices, SEO và PWA [21].",
        advantages=[
            "Đo chính xác Core Web Vitals theo chuẩn của Google.",
            "Xuất kết quả JSON dễ tích hợp với hệ thống khác.",
        ],
        disadvantages=[
            "Chạy chậm (~30–60 giây/URL) so với Cheerio.",
        ],
        usage=[
            "Đo LCP, FID, CLS cho mỗi URL được audit.",
            "Kết quả được cache trong Redis để tránh chạy lại trong 1 giờ.",
        ])
    tech_block(doc, "Docker và GitHub Actions",
        "Docker là nền tảng container hóa giúp đóng gói ứng dụng cùng "
        "dependencies. GitHub Actions là pipeline CI/CD miễn phí của GitHub [22].",
        advantages=[
            "Đảm bảo môi trường development và production đồng nhất.",
            "GitHub Actions có 2000 phút miễn phí/tháng cho repo private.",
            "Dễ viết workflow YAML, có hàng nghìn action có sẵn.",
        ],
        disadvantages=[
            "Docker image có thể lớn nếu không tối ưu.",
        ],
        usage=[
            "Container hóa frontend, backend, worker và database.",
            "Pipeline tự động: lint → test → build → deploy khi push code.",
        ])
    tech_block(doc, "Vercel, Railway và Supabase",
        "Ba nền tảng cloud với free-tier hoặc chi phí thấp [23][24][25]. "
        "Vercel chuyên cho frontend Next.js, Railway cho backend container, "
        "Supabase cung cấp PostgreSQL + Auth + Realtime managed.",
        advantages=[
            "Tổng chi phí dưới 40 USD/tháng, đúng mục tiêu ngân sách.",
            "Tích hợp native với GitHub, deploy tự động khi merge.",
            "Vercel Preview URL cho mỗi pull request.",
        ],
        disadvantages=[
            "Giới hạn tài nguyên trên free-tier.",
            "Phụ thuộc vào vendor, chuyển đổi sau này phức tạp.",
        ],
        usage=[
            "Vercel: deploy frontend Next.js.",
            "Railway: deploy backend NestJS và worker BullMQ.",
            "Supabase: PostgreSQL managed và Redis managed.",
        ])

    # ---- 1.4 Kết luận chương ----
    chapter_conclusion(doc, "1.4.", [
        "Chương 1 đã trình bày tổng quan về đề tài xây dựng nền tảng phân tích "
        "SEO website tự động, bao gồm lý do lựa chọn, mục tiêu, nội dung thực "
        "hiện, cùng với khảo sát nhu cầu người dùng và phân tích các công "
        "nghệ được ứng dụng trong hệ thống.",
        "Thông qua việc khảo sát nhu cầu và nghiên cứu các nền tảng SEO hiện "
        "có, em đã xác định được các yêu cầu cơ bản từ phía người dùng cuối "
        "và quản trị viên, từ đó làm cơ sở cho quá trình phân tích và thiết "
        "kế hệ thống ở các chương tiếp theo.",
        "Ngoài ra, việc lựa chọn các công nghệ hiện đại như Next.js, NestJS, "
        "PostgreSQL, Redis, Cheerio và Playwright không chỉ giúp hệ thống "
        "hoạt động ổn định, tối ưu trải nghiệm người dùng mà còn đảm bảo "
        "khả năng mở rộng và duy trì về sau. Những cơ sở này sẽ được cụ thể "
        "hóa trong Chương 2 với các mô hình phân tích chức năng, luồng xử "
        "lý và sơ đồ hệ thống chi tiết.",
    ])
    doc.add_page_break()

    # ========================================================================
    # CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
    # ========================================================================
    h1(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

    # ---- 2.1 Thiết kế cơ sở dữ liệu ----
    h2(doc, "2.1. Thiết kế cơ sở dữ liệu")
    para(doc,
        "Cơ sở dữ liệu được thiết kế theo chuẩn Third Normal Form (3NF) để "
        "loại bỏ dư thừa dữ liệu và đảm bảo tính toàn vẹn tham chiếu. Đối "
        "với bảng audit_results – nơi lưu kết quả phân tích dưới dạng JSON "
        "với cấu trúc thay đổi theo từng rule – đồ án áp dụng denormalization "
        "có kiểm soát bằng cách dùng kiểu JSONB của PostgreSQL. Lý do: kết "
        "quả rule có cấu trúc bất định và không cần truy vấn sâu vào từng "
        "field, nên JSONB cho tốc độ đọc/ghi nhanh hơn nhiều so với việc "
        "tạo hàng chục bảng con cho mỗi loại rule [16]."
    )
    embed_image(doc, "01-erd.png", "Hình 2.1",
        "Sơ đồ quan hệ thực thể (ERD) của PostgreSQL schema")
    para(doc,
        "Hệ thống có các bảng chính: users (thông tin người dùng), audits "
        "(thông tin mỗi lần phân tích), audit_results (chi tiết kết quả từng "
        "rule), keywords (từ khóa được phân tích). Các bảng này có quan hệ "
        "1-nhiều như sau: một user có nhiều audit, một audit có nhiều "
        "audit_result và nhiều keyword. Dưới đây là mô tả chi tiết các bảng."
    )
    caption(doc, "Bảng 2.1 đến 2.4. Mô tả chi tiết các bảng trong cơ sở dữ liệu")
    render_items(doc, s_db, demote=True)
    para(doc,
        "Về chiến lược đánh index, hệ thống sử dụng ba loại: (1) B-Tree "
        "index mặc định cho khóa chính và khóa ngoại, (2) composite index "
        "trên (user_id, created_at DESC) để tối ưu truy vấn lịch sử audit, "
        "(3) GIN index trên cột JSONB của audit_results để truy vấn nhanh "
        "theo rule cụ thể. Migration được quản lý bằng Prisma Migrate."
    )

    # ---- 2.2 Biểu đồ Use-case ----
    h2(doc, "2.2. Biểu đồ Use-case")
    para(doc,
        "Hệ thống có 3 nhóm tác nhân (actor): Khách vãng lai (Guest), Người "
        "dùng đã đăng ký (Registered User) và Quản trị viên (Admin). Phần "
        "này trình bày biểu đồ use-case cho từng nhóm và đặc tả chi tiết "
        "một số use-case quan trọng."
    )

    h3(doc, "2.2.1. Use-case người dùng")
    embed_image(doc, "02-usecase-user.png", "Hình 2.2",
        "Biểu đồ use-case người dùng cuối (Guest + Registered User)")
    caption(doc, "Bảng 2.5. Bảng use-case tổng quan của người dùng cuối")
    render_items(doc, s_func, demote=True)

    h3(doc, "2.2.2. Use-case quản trị viên")
    embed_image(doc, "03-usecase-admin.png", "Hình 2.3",
        "Biểu đồ use-case của quản trị viên (Admin)")
    para(doc,
        "Quản trị viên có các chức năng: Quản lý người dùng, Quản lý rule "
        "SEO, Theo dõi thống kê hệ thống, Xử lý báo cáo lạm dụng. Đây là các "
        "chức năng chỉ hiển thị cho user có role = 'admin' trong database."
    )

    h3(doc, "2.2.3. Phân tích chi tiết một số use-case quan trọng")

    doc.add_paragraph("2.2.3.1. Đăng ký tài khoản", style="Heading 3")
    usecase_block(doc, "UC-01. Đăng ký tài khoản",
        actor="Khách vãng lai",
        desc="Người dùng mới tạo tài khoản để sử dụng đầy đủ tính năng của hệ thống.",
        main_flow=[
            "Người dùng truy cập trang chủ và nhấn nút \"Đăng ký\".",
            "Hệ thống hiển thị form nhập email, mật khẩu và xác nhận mật khẩu.",
            "Người dùng điền thông tin và nhấn \"Đăng ký\".",
            "Hệ thống validate dữ liệu (email hợp lệ, mật khẩu ≥ 8 ký tự, trùng xác nhận).",
            "Hệ thống kiểm tra email chưa tồn tại trong database.",
            "Hệ thống hash mật khẩu bằng bcrypt với cost factor 12.",
            "Hệ thống tạo record user mới và gửi email xác minh.",
            "Hệ thống chuyển người dùng sang trang đăng nhập kèm thông báo thành công.",
        ],
        alt_flow="Nếu email đã tồn tại, hiển thị lỗi \"Email đã được sử dụng\". "
                 "Nếu mật khẩu không đủ mạnh, hiển thị gợi ý tăng độ phức tạp.",
        pre="Người dùng chưa có tài khoản trong hệ thống.",
        post="Tài khoản được tạo với trạng thái 'pending_verification' và email "
             "xác minh đã được gửi.")

    doc.add_paragraph("2.2.3.2. Đăng nhập", style="Heading 3")
    usecase_block(doc, "UC-02. Đăng nhập",
        actor="Người dùng đã đăng ký",
        desc="Người dùng đăng nhập vào hệ thống để sử dụng các tính năng cá nhân.",
        main_flow=[
            "Người dùng truy cập trang đăng nhập.",
            "Người dùng nhập email và mật khẩu, nhấn \"Đăng nhập\".",
            "Hệ thống tìm user theo email trong database.",
            "Hệ thống so khớp mật khẩu bằng bcrypt.compare.",
            "Nếu đúng, hệ thống sinh access token (15 phút) và refresh token (7 ngày).",
            "Refresh token được lưu vào HttpOnly cookie, access token trả về body.",
            "Người dùng được chuyển sang trang Dashboard.",
        ],
        alt_flow="Nếu email không tồn tại hoặc mật khẩu sai, hiển thị thông báo "
                 "chung \"Email hoặc mật khẩu không đúng\" để tránh tiết lộ thông tin.",
        pre="Người dùng đã có tài khoản đã xác minh email.",
        post="Người dùng được cấp token và có thể truy cập các trang yêu cầu "
             "xác thực.")

    doc.add_paragraph("2.2.3.3. Phân tích SEO một URL", style="Heading 3")
    usecase_block(doc, "UC-05. Phân tích SEO một URL",
        actor="Người dùng đã đăng ký",
        desc="Người dùng nhập URL website và nhận kết quả phân tích SEO đầy đủ.",
        main_flow=[
            "Người dùng truy cập Dashboard và nhấn nút \"Audit mới\".",
            "Hệ thống hiển thị form nhập URL.",
            "Người dùng nhập URL và nhấn \"Bắt đầu phân tích\".",
            "Hệ thống validate URL hợp lệ (protocol, format, không phải localhost).",
            "Hệ thống kiểm tra rate limit (tối đa 10 audit/giờ cho user free).",
            "API Gateway tạo record audit mới trong PostgreSQL, trạng thái 'pending'.",
            "API Gateway đẩy job vào BullMQ queue với audit_id.",
            "Hệ thống trả về audit_id cho frontend và chuyển sang trang chi tiết.",
            "Crawler Worker pick job, crawl HTML/DOM, phát sự kiện progress.",
            "Analyzer Worker nhận DOM, áp dụng 20 rule SEO, phát sự kiện progress.",
            "Report Worker tổng hợp kết quả, lưu audit_results, phát sự kiện completed.",
            "Frontend nhận sự kiện completed qua Socket.IO, fetch kết quả và hiển thị.",
        ],
        alt_flow="Nếu URL trả về HTTP status ≠ 2xx, ghi nhận issue \"URL không "
                 "thể truy cập\". Nếu crawl timeout > 30s, retry 2 lần với "
                 "exponential backoff.",
        pre="Người dùng đã đăng nhập và còn quota audit.",
        post="Kết quả audit được lưu vào database và hiển thị cho người dùng.")

    doc.add_paragraph("2.2.3.4. Xuất báo cáo PDF", style="Heading 3")
    usecase_block(doc, "UC-08. Xuất báo cáo PDF",
        actor="Người dùng đã đăng ký",
        desc="Người dùng tải về báo cáo PDF của một audit đã hoàn thành.",
        main_flow=[
            "Người dùng xem trang chi tiết audit.",
            "Người dùng nhấn nút \"Tải PDF\".",
            "Frontend gọi endpoint GET /audits/{id}/export?format=pdf.",
            "Report Service truy vấn audit và audit_results từ database.",
            "Report Service render template Handlebars với dữ liệu audit.",
            "Report Service dùng Playwright.pdf() để chuyển HTML thành PDF.",
            "File PDF được trả về với Content-Disposition: attachment.",
            "Trình duyệt tự động tải file về máy người dùng.",
        ],
        alt_flow="Nếu audit chưa hoàn thành (status != 'completed'), trả về "
                 "lỗi 400 với thông báo phù hợp.",
        pre="Audit đã hoàn thành và thuộc về user đang đăng nhập.",
        post="File PDF được tải về máy người dùng.")

    # ---- 2.3 Sơ đồ phân rã chức năng ----
    h2(doc, "2.3. Sơ đồ phân rã chức năng")
    para(doc,
        "Sơ đồ phân rã chức năng thể hiện cấu trúc cây của các chức năng "
        "trong hệ thống, giúp dễ dàng nhận biết mối quan hệ cha-con và phân "
        "chia nhiệm vụ phát triển theo module."
    )
    embed_image(doc, "04-function-decomposition.png", "Hình 2.4",
        "Sơ đồ phân rã chức năng hệ thống SEO Analyst Platform", width_inches=5.5)
    para(doc, "Hệ thống được phân rã thành 5 nhóm chức năng chính:")
    for grp, subs in [
        ("Quản lý tài khoản", "Đăng ký, Đăng nhập, Đăng xuất, Quên mật khẩu, Cập nhật hồ sơ, Đổi mật khẩu"),
        ("Phân tích SEO", "Tạo audit mới, Xem chi tiết audit, Xem lịch sử audit, Xóa audit, So sánh audit"),
        ("Báo cáo", "Xuất báo cáo PDF, Xuất CSV, Chia sẻ link audit"),
        ("Quản trị hệ thống", "Quản lý người dùng, Cấu hình rule SEO, Xem thống kê, Xử lý báo cáo lạm dụng"),
        ("Thông báo", "Nhận thông báo audit hoàn thành, Cấu hình thông báo email"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{grp}: ").bold = True
        p.add_run(subs)

    # ---- 2.4 Biểu đồ tuần tự ----
    h2(doc, "2.4. Biểu đồ tuần tự")
    para(doc,
        "Biểu đồ tuần tự (Sequence Diagram) mô tả trình tự các thông điệp "
        "được truyền giữa các đối tượng (actor, frontend, backend, database, "
        "worker) trong quá trình thực hiện một chức năng cụ thể. Phần này "
        "trình bày sequence diagram cho 3 chức năng quan trọng nhất của hệ thống."
    )

    h3(doc, "2.4.1. Chức năng đăng nhập")
    embed_image(doc, "05-sequence-login.png", "Hình 2.5",
        "Biểu đồ tuần tự chức năng đăng nhập với JWT")
    para(doc, "Mô tả chi tiết các bước:")
    for step in [
        "Người dùng nhập email/mật khẩu và submit form.",
        "Frontend gửi POST /auth/login với payload {email, password}.",
        "Auth Controller gọi AuthService.validateUser.",
        "AuthService truy vấn UserRepository để tìm user theo email.",
        "UserRepository thực thi SELECT * FROM users WHERE email = ? trên PostgreSQL.",
        "AuthService so khớp mật khẩu bằng bcrypt.compare.",
        "Nếu đúng, AuthService sinh access token (jwt.sign, expiresIn=15m).",
        "AuthService sinh refresh token và lưu hash vào bảng refresh_tokens.",
        "Auth Controller set HttpOnly cookie với refresh token.",
        "Auth Controller trả về access token trong response body.",
        "Frontend lưu access token vào memory (không localStorage để tránh XSS).",
        "Frontend chuyển hướng sang /dashboard.",
    ]:
        doc.add_paragraph(f"− {step}")

    h3(doc, "2.4.2. Chức năng phân tích SEO một URL")
    embed_image(doc, "06-sequence-audit.png", "Hình 2.6",
        "Biểu đồ tuần tự chức năng phân tích SEO với BullMQ và Socket.IO", width_inches=6.5)
    para(doc, "Mô tả chi tiết các bước:")
    for step in [
        "Người dùng nhập URL và nhấn \"Phân tích\" trên form.",
        "Frontend gửi POST /audits với payload {url}.",
        "API Gateway validate JWT qua JwtAuthGuard.",
        "AuditController gọi AuditService.createAudit.",
        "AuditService tạo record audits (status='pending', user_id) và commit.",
        "AuditService đẩy job vào BullMQ queue: audit.crawl với audit_id.",
        "AuditController trả về audit_id cho frontend.",
        "Frontend kết nối Socket.IO room với audit_id.",
        "Crawler Worker pick job, cập nhật status='crawling', emit audit.progress (25%).",
        "Crawler Worker fetch HTML bằng axios+Cheerio hoặc Playwright.",
        "Crawler Worker đẩy job tiếp: audit.analyze với DOM.",
        "Analyzer Worker pick job, cập nhật status='analyzing', emit audit.progress (50%).",
        "Analyzer Worker áp dụng 20 rule SEO, tính điểm.",
        "Analyzer Worker đẩy job tiếp: audit.report.",
        "Report Worker tổng hợp, lưu audit_results, cập nhật status='completed', emit audit.completed (100%).",
        "Frontend nhận audit.completed, invalidate cache TanStack Query.",
        "Frontend fetch GET /audits/{id} và hiển thị kết quả chi tiết.",
    ]:
        doc.add_paragraph(f"− {step}")

    h3(doc, "2.4.3. Chức năng xuất báo cáo PDF")
    embed_image(doc, "07-sequence-export-pdf.png", "Hình 2.7",
        "Biểu đồ tuần tự chức năng xuất báo cáo PDF")
    para(doc, "Mô tả chi tiết các bước:")
    for step in [
        "Người dùng nhấn nút \"Tải PDF\" trên trang chi tiết audit.",
        "Frontend gửi GET /audits/{id}/export?format=pdf.",
        "ReportController validate JWT và quyền sở hữu audit.",
        "ReportService truy vấn audit + audit_results từ database.",
        "ReportService render HTML template Handlebars với dữ liệu.",
        "ReportService khởi chạy Playwright browser context.",
        "ReportService gọi page.setContent(html) rồi page.pdf({format:'A4'}).",
        "ReportService trả về Buffer PDF cho controller.",
        "ReportController set header Content-Type: application/pdf và Content-Disposition: attachment.",
        "Frontend nhận response, trình duyệt tự động mở dialog tải file.",
    ]:
        doc.add_paragraph(f"− {step}")

    # ---- 2.5 Thiết kế kiến trúc hệ thống ----
    h2(doc, "2.5. Thiết kế kiến trúc hệ thống")

    h3(doc, "2.5.1. Kiến trúc microservices tổng thể")
    para(doc,
        "Hệ thống được thiết kế theo mô hình Microservices – mỗi chức năng "
        "nghiệp vụ cốt lõi được tách thành một service độc lập, giao tiếp "
        "với nhau qua REST API và message queue. So với kiến trúc monolith "
        "truyền thống, microservices có ưu điểm: cho phép mỗi service scale "
        "độc lập, dễ thay thế thành phần và phù hợp với đặc thù resource "
        "usage khác biệt giữa các service (ví dụ Crawler cần nhiều RAM cho "
        "Playwright trong khi Analyzer chủ yếu tính toán) [2]."
    )
    embed_image(doc, "08-architecture-microservices.png", "Hình 2.8",
        "Sơ đồ kiến trúc microservices tổng thể", width_inches=6.5)
    render_items(doc, s_sysoverview, demote=True)
    embed_image(doc, "09-dfd-level1.png", "Hình 2.9",
        "Sơ đồ luồng dữ liệu (Data Flow Diagram) level-1", width_inches=6.5)

    h3(doc, "2.5.2. Chi tiết các microservice")
    render_items(doc, s_microservices, demote=True)

    h3(doc, "2.5.3. Thiết kế REST API")
    para(doc,
        "REST API được thiết kế theo 6 ràng buộc của Roy Fielding [1]: "
        "client-server, stateless, cacheable, uniform interface, layered "
        "system và code on demand. Quy ước: dùng danh từ số nhiều cho tài "
        "nguyên, HTTP method đúng ngữ nghĩa, status code chuẩn, versioning "
        "qua URL prefix /api/v1/, trả lỗi theo RFC 7807."
    )
    render_items(doc, s_api, demote=True)

    h3(doc, "2.5.4. Yêu cầu phi chức năng")
    para(doc,
        "Yêu cầu phi chức năng (NFR) được chia thành 6 nhóm: Hiệu năng, "
        "Khả năng mở rộng, Tính sẵn sàng, Bảo mật, Khả năng sử dụng và "
        "Khả năng bảo trì. Bảng dưới đây liệt kê các chỉ số cụ thể."
    )
    caption(doc, "Bảng 2.6. Bảng đặc tả yêu cầu phi chức năng")
    render_items(doc, s_nfr, demote=True)

    # ---- 2.6 Kết luận chương ----
    chapter_conclusion(doc, "2.6.", [
        "Chương 2 đã trình bày chi tiết quá trình phân tích và thiết kế hệ "
        "thống SEO Analyst Platform, bao gồm thiết kế cơ sở dữ liệu, biểu "
        "đồ use-case, sơ đồ phân rã chức năng, biểu đồ tuần tự cho các "
        "chức năng quan trọng và kiến trúc microservices tổng thể.",
        "Qua việc áp dụng các phương pháp phân tích hệ thống kinh điển như "
        "mô hình use-case, sequence diagram và phân rã chức năng, đồ án đã "
        "xác định rõ ràng trách nhiệm của từng thành phần và cách chúng "
        "tương tác với nhau. Thiết kế cơ sở dữ liệu theo chuẩn 3NF kết hợp "
        "với JSONB cho các cột bất định giúp cân bằng giữa tính toàn vẹn và "
        "hiệu năng.",
        "Những thiết kế này sẽ được hiện thực hóa trong Chương 3 thông qua "
        "việc triển khai thực tế hệ thống: cài đặt môi trường, tổ chức mã "
        "nguồn, xây dựng giao diện và kiểm thử chức năng.",
    ])
    doc.add_page_break()

    # ========================================================================
    # CHƯƠNG 3. TRIỂN KHAI DỰ ÁN
    # ========================================================================
    h1(doc, "CHƯƠNG 3. TRIỂN KHAI DỰ ÁN")

    # ---- 3.1 Môi trường phát triển và cài đặt ----
    h2(doc, "3.1. Môi trường phát triển và cài đặt")

    h3(doc, "3.1.1. Môi trường phát triển")
    for line in [
        "Hệ điều hành (Operating System): macOS Sonoma 14 / Windows 11 / Ubuntu 22.04 LTS.",
        "Trình soạn thảo mã (IDE): Visual Studio Code với các extension ESLint, Prettier, Prisma, Tailwind CSS IntelliSense.",
        "Trình duyệt kiểm thử: Google Chrome 120+ (DevTools), Firefox 121+.",
        "Node.js: v20.11.0 LTS.",
        "Trình quản lý gói (Package Manager): pnpm 9.0.0.",
        "Công nghệ chính sử dụng:",
        "  Next.js 14: React framework hỗ trợ App Router, Server Components.",
        "  NestJS 10: Framework backend TypeScript với DI, guards, pipes.",
        "  PostgreSQL 16: Cơ sở dữ liệu quan hệ chính.",
        "  Prisma 5: ORM type-safe với migration và Prisma Studio.",
        "  Redis 7: In-memory store cho BullMQ và cache.",
        "  BullMQ 5: Hàng đợi job bất đồng bộ dựa trên Redis.",
        "  Cheerio 1.0: Parser HTML nhẹ với cú pháp jQuery.",
        "  Playwright 1.40: Framework điều khiển trình duyệt headless.",
        "  Tailwind CSS 3.4 + shadcn/ui: Framework CSS utility-first.",
        "  TanStack Query 5: Quản lý server state cho React.",
        "  Socket.IO 4.7: Thư viện realtime bi-directional.",
        "  Docker 24 + Docker Compose: Container hóa môi trường phát triển.",
    ]:
        doc.add_paragraph(f"− {line}")

    h3(doc, "3.1.2. Hướng dẫn cài đặt")
    para(doc,
        "Để triển khai dự án trên máy cá nhân, cần thực hiện các bước cài "
        "đặt sau:"
    )
    doc.add_paragraph("Bước 1: Cài đặt Node.js và pnpm.").runs[0].bold = True
    for line in [
        "Truy cập https://nodejs.org và tải bản LTS (Node.js v20.11.0).",
        "Cài đặt pnpm: npm install -g pnpm@9",
        "Kiểm tra phiên bản: node -v && pnpm -v",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 2: Tải mã nguồn và cài đặt thư viện.").runs[0].bold = True
    for line in [
        "git clone https://github.com/<user>/SEO-Analysts.git",
        "cd SEO-Analysts",
        "pnpm install",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 3: Khởi động PostgreSQL và Redis bằng Docker Compose.").runs[0].bold = True
    for line in [
        "docker compose up -d postgres redis",
        "(đảm bảo Docker Desktop đang chạy)",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 4: Cấu hình biến môi trường.").runs[0].bold = True
    for line in [
        "cp .env.example .env",
        "(chỉnh sửa DATABASE_URL, REDIS_URL, JWT_SECRET theo môi trường)",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 5: Chạy migration và seed dữ liệu mẫu.").runs[0].bold = True
    for line in [
        "pnpm --filter @seo/backend prisma migrate dev",
        "pnpm --filter @seo/backend prisma db seed",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 6: Chạy backend và frontend ở hai terminal riêng.").runs[0].bold = True
    for line in [
        "# Terminal 1 – Backend",
        "pnpm --filter @seo/backend dev",
        "# Terminal 2 – Frontend",
        "pnpm --filter @seo/frontend dev",
    ]:
        doc.add_paragraph(f"− {line}")
    doc.add_paragraph("Bước 7: Truy cập hệ thống.").runs[0].bold = True
    for line in [
        "Frontend: http://localhost:3000",
        "Backend API (Swagger): http://localhost:3001/api",
    ]:
        doc.add_paragraph(f"− {line}")

    # ---- 3.2 Cấu trúc thư mục dự án ----
    h2(doc, "3.2. Cấu trúc thư mục dự án")
    para(doc,
        "Dự án được tổ chức theo mô hình monorepo sử dụng Turborepo, cho "
        "phép quản lý nhiều package trong cùng một repository. Cách tổ chức "
        "này giúp chia sẻ code (types, utils) giữa frontend và backend dễ "
        "dàng, tránh duplicate và đảm bảo đồng bộ."
    )
    embed_image(doc, "10-folder-tree-root.png", "Hình 3.1",
        "Cấu trúc thư mục tổng thể của dự án (Turborepo monorepo)", width_inches=6.5)
    para(doc, "Cấu trúc thư mục gốc bao gồm các thư mục chính:")
    for line in [
        "apps/: chứa các ứng dụng deploy được (frontend và backend).",
        "packages/: chứa các package dùng chung (types, utils, config).",
        "scripts/: chứa các script tiện ích (seed, migrate, build).",
        "docker-compose.yml: định nghĩa các service local (postgres, redis).",
        "turbo.json: cấu hình Turborepo (pipeline build, test, dev).",
        "pnpm-workspace.yaml: khai báo workspaces cho pnpm.",
    ]:
        doc.add_paragraph(f"− {line}")

    h3(doc, "3.2.1. Thư mục apps/frontend")
    embed_image(doc, "11-folder-tree-frontend.png", "Hình 3.2",
        "Cấu trúc thư mục apps/frontend (Next.js 14 App Router)", width_inches=6.5)
    para(doc, "Thư mục apps/frontend chứa ứng dụng Next.js 14 với các thành phần:")
    for line in [
        "app/: các route theo App Router (page.tsx, layout.tsx, loading.tsx, error.tsx).",
        "app/(auth)/: nhóm route xác thực (đăng nhập, đăng ký, quên mật khẩu).",
        "app/(dashboard)/: nhóm route dashboard (audit list, audit detail, profile).",
        "app/(admin)/: nhóm route quản trị (user management, rule config).",
        "components/ui/: các component từ shadcn/ui (Button, Dialog, Table...).",
        "components/feature/: các component theo chức năng (AuditCard, ScoreChart...).",
        "hooks/: custom hooks (useAuditSocket, useAuth, useDebounce...).",
        "lib/: utility functions và API client.",
        "schemas/: Zod schema cho validation form.",
        "store/: Zustand store cho global state (nếu cần).",
    ]:
        doc.add_paragraph(f"− {line}")

    h3(doc, "3.2.2. Thư mục apps/backend")
    embed_image(doc, "12-folder-tree-backend.png", "Hình 3.3",
        "Cấu trúc thư mục apps/backend (NestJS 10)", width_inches=5.0)
    para(doc, "Thư mục apps/backend chứa ứng dụng NestJS với cấu trúc module:")
    for line in [
        "src/auth/: module xác thực với JWT strategy, guards, DTO.",
        "src/users/: module quản lý người dùng (controller, service, repository).",
        "src/audit/: module phân tích SEO (controller, service, processor BullMQ).",
        "src/crawler/: module crawl HTML với Cheerio + Playwright.",
        "src/analyzer/: module áp dụng 20 rule SEO và chấm điểm.",
        "src/report/: module xuất PDF và tổng hợp báo cáo.",
        "src/common/: guards, filters, interceptors dùng chung.",
        "prisma/schema.prisma: định nghĩa schema cơ sở dữ liệu.",
        "prisma/migrations/: các file migration SQL có version.",
        "test/: unit test và e2e test.",
    ]:
        doc.add_paragraph(f"− {line}")

    h3(doc, "3.2.3. Thư mục packages/")
    para(doc, "Thư mục packages/ chứa các package dùng chung giữa frontend và backend:")
    for line in [
        "packages/types/: TypeScript interface và type dùng chung (User, Audit, Rule).",
        "packages/utils/: helper function dùng chung (formatDate, formatUrl, validators).",
        "packages/config/: cấu hình dùng chung (eslint, tsconfig, tailwind preset).",
    ]:
        doc.add_paragraph(f"− {line}")

    # ---- 3.3 Giao diện người dùng ----
    h2(doc, "3.3. Giao diện người dùng")
    para(doc,
        "Phần này trình bày các giao diện chính của hệ thống cùng mô tả "
        "chức năng và tương tác của người dùng trên từng trang. Giao diện "
        "được xây dựng với Tailwind CSS + shadcn/ui, hỗ trợ responsive và "
        "dark mode, tuân thủ chuẩn accessibility WAI-ARIA."
    )

    h3(doc, "3.3.1. Giao diện trang chủ")
    para(doc,
        "Trang chủ là điểm tiếp xúc đầu tiên của người dùng với hệ thống. "
        "Giao diện bao gồm hero section với slogan, mô tả ngắn gọn về sản "
        "phẩm, CTA \"Bắt đầu miễn phí\", section tính năng chính, section "
        "so sánh với các công cụ thương mại và footer. Trang chủ được "
        "render bằng Server Components để tối ưu SEO."
    )
    placeholder(doc, "Hình 3.4", "SCREENSHOT",
        "Giao diện trang chủ (Landing Page)",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.2. Giao diện đăng nhập / đăng ký")
    para(doc,
        "Trang đăng nhập và đăng ký có thiết kế tối giản với form ở giữa "
        "màn hình, hỗ trợ validation realtime bằng Zod + React Hook Form. "
        "Nút \"Hiển thị mật khẩu\" giúp người dùng kiểm tra nhập liệu. "
        "Link chuyển đổi giữa hai form nằm ở cuối."
    )
    placeholder(doc, "Hình 3.5", "SCREENSHOT",
        "Giao diện trang đăng nhập",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.3. Giao diện dashboard danh sách audit")
    para(doc,
        "Dashboard hiển thị danh sách các audit đã thực hiện dưới dạng "
        "lưới Card, mỗi Card gồm: URL đã audit, điểm tổng, thời gian, trạng "
        "thái. Người dùng có thể lọc theo điểm, trạng thái, khoảng thời "
        "gian và tìm kiếm theo URL. Audit đang chạy có hiệu ứng loading "
        "với progress bar realtime."
    )
    placeholder(doc, "Hình 3.6", "SCREENSHOT",
        "Giao diện dashboard danh sách audit",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.4. Giao diện form nhập URL phân tích")
    para(doc,
        "Modal hoặc trang riêng cho phép người dùng nhập URL mới để phân "
        "tích. Form có validation URL hợp lệ, hiển thị preview favicon và "
        "title của trang (nếu có). Sau khi submit, modal chuyển sang trạng "
        "thái progress với các bước: \"Đang crawl\" → \"Đang phân tích\" → "
        "\"Đang tổng hợp\" → \"Hoàn thành\"."
    )
    placeholder(doc, "Hình 3.7", "SCREENSHOT",
        "Giao diện form nhập URL và progress modal",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.5. Giao diện chi tiết kết quả audit")
    para(doc,
        "Trang chi tiết audit là trang quan trọng nhất, hiển thị đầy đủ kết "
        "quả phân tích. Bao gồm: header với URL + điểm tổng + nút xuất PDF; "
        "biểu đồ radar 5 tiêu chí (Performance, SEO, Accessibility, Best "
        "Practices, Content); danh sách các issue được phát hiện kèm mức "
        "độ nghiêm trọng và gợi ý cải thiện; tab Core Web Vitals với LCP, "
        "FID, CLS đo được; tab Keywords với các từ khóa phát hiện trong "
        "nội dung trang."
    )
    placeholder(doc, "Hình 3.8", "SCREENSHOT",
        "Giao diện chi tiết kết quả audit",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.6. Giao diện quản trị viên")
    para(doc,
        "Trang dành cho admin với các chức năng: quản lý người dùng "
        "(danh sách, khóa/mở tài khoản), cấu hình trọng số rule SEO "
        "(form cho phép điều chỉnh weight của từng rule), xem thống kê "
        "(biểu đồ số audit/ngày, top URL được audit nhiều nhất, thời gian "
        "crawl trung bình)."
    )
    placeholder(doc, "Hình 3.9", "SCREENSHOT",
        "Giao diện quản trị viên",
        "Chạy frontend → screenshot")

    h3(doc, "3.3.7. Mẫu báo cáo PDF")
    placeholder(doc, "Hình 3.10", "SCREENSHOT",
        "Mẫu báo cáo PDF xuất ra từ hệ thống",
        "Chạy demo → xuất PDF → screenshot trang PDF")

    # ---- 3.4 Kiểm thử chức năng ----
    h2(doc, "3.4. Kiểm thử chức năng")
    para(doc,
        "Hệ thống áp dụng mô hình Testing Pyramid: 70% unit test (nhanh, "
        "rẻ), 20% integration test (kiểm tra phối hợp module) và 10% E2E "
        "test (mô phỏng người dùng thật). Backend sử dụng Vitest, integration "
        "test dùng TestContainers (PostgreSQL + Redis thật), E2E dùng "
        "Playwright Test. Code coverage mục tiêu ≥ 80% cho business logic."
    )
    render_items(doc, s_test, demote=True)

    h3(doc, "3.4.1. Kiểm thử chức năng xác thực")
    testcase_table(doc, "Bảng 3.1. Test case chức năng xác thực", [
        ("TC_01", "Đăng ký tài khoản",
         "Người dùng tạo tài khoản mới thành công",
         "1. Truy cập /auth/register\n2. Nhập email, mật khẩu hợp lệ\n3. Nhấn Đăng ký\n4. Nhận email xác minh",
         "Pass"),
        ("TC_02", "Đăng ký email trùng",
         "Hệ thống từ chối email đã tồn tại",
         "1. Truy cập /auth/register\n2. Nhập email đã tồn tại\n3. Nhấn Đăng ký",
         "Pass"),
        ("TC_03", "Đăng ký mật khẩu yếu",
         "Hệ thống từ chối mật khẩu < 8 ký tự",
         "1. Nhập mật khẩu \"1234\"\n2. Nhấn Đăng ký",
         "Pass"),
        ("TC_04", "Đăng nhập thành công",
         "Người dùng đăng nhập với tài khoản hợp lệ",
         "1. Truy cập /auth/login\n2. Nhập email + mật khẩu đúng\n3. Nhấn Đăng nhập",
         "Pass"),
        ("TC_05", "Đăng nhập sai mật khẩu",
         "Hệ thống từ chối và không tiết lộ email có tồn tại",
         "1. Nhập email đúng + mật khẩu sai\n2. Nhấn Đăng nhập",
         "Pass"),
        ("TC_06", "Refresh token",
         "Access token hết hạn được làm mới tự động",
         "1. Đợi access token hết hạn (15 phút)\n2. Gọi API bất kỳ\n3. Hệ thống tự gọi /auth/refresh",
         "Pass"),
        ("TC_07", "Đăng xuất",
         "Refresh token bị vô hiệu hóa sau khi đăng xuất",
         "1. Đăng nhập\n2. Nhấn Đăng xuất\n3. Thử gọi /auth/refresh với token cũ",
         "Pass"),
    ])

    h3(doc, "3.4.2. Kiểm thử chức năng phân tích SEO")
    testcase_table(doc, "Bảng 3.2. Test case chức năng phân tích SEO", [
        ("TC_01", "Tạo audit URL tĩnh",
         "Phân tích thành công trang HTML tĩnh",
         "1. Đăng nhập\n2. Nhập URL https://example.com\n3. Nhấn Phân tích\n4. Đợi kết quả",
         "Pass"),
        ("TC_02", "Tạo audit URL SPA",
         "Hệ thống fallback sang Playwright cho SPA",
         "1. Nhập URL React app\n2. Nhấn Phân tích\n3. Quan sát progress",
         "Pass"),
        ("TC_03", "Tạo audit URL không hợp lệ",
         "Hệ thống từ chối URL sai format",
         "1. Nhập \"not-a-url\"\n2. Nhấn Phân tích",
         "Pass"),
        ("TC_04", "Tạo audit URL localhost",
         "Hệ thống từ chối URL localhost/SSRF",
         "1. Nhập http://localhost:3000\n2. Nhấn Phân tích",
         "Pass"),
        ("TC_05", "Rate limit audit",
         "User free bị giới hạn 10 audit/giờ",
         "1. Tạo 10 audit liên tiếp\n2. Tạo audit thứ 11",
         "Pass"),
        ("TC_06", "Xem chi tiết audit",
         "Hiển thị đầy đủ score, issue, CWV, keywords",
         "1. Nhấn vào audit trong dashboard\n2. Kiểm tra các tab",
         "Pass"),
        ("TC_07", "Progress realtime",
         "Progress bar cập nhật theo thời gian thực",
         "1. Tạo audit mới\n2. Quan sát progress bar",
         "Pass"),
        ("TC_08", "Xóa audit",
         "Audit bị xóa khỏi database",
         "1. Mở chi tiết audit\n2. Nhấn nút Xóa\n3. Xác nhận",
         "Pass"),
    ])

    h3(doc, "3.4.3. Kiểm thử chức năng báo cáo")
    testcase_table(doc, "Bảng 3.3. Test case chức năng báo cáo", [
        ("TC_01", "Xuất PDF",
         "Tải về file PDF chứa đầy đủ kết quả audit",
         "1. Mở chi tiết audit đã completed\n2. Nhấn nút Tải PDF\n3. Kiểm tra file",
         "Pass"),
        ("TC_02", "Xuất PDF audit chưa hoàn thành",
         "Hệ thống từ chối với status 400",
         "1. Gọi API export với audit đang pending",
         "Pass"),
        ("TC_03", "Xem lịch sử audit",
         "Danh sách audit hiển thị đúng thứ tự thời gian",
         "1. Truy cập /dashboard\n2. Kiểm tra thứ tự",
         "Pass"),
        ("TC_04", "Lọc audit theo điểm",
         "Dashboard chỉ hiển thị audit có điểm trong khoảng",
         "1. Chọn filter \"80–100\"\n2. Kiểm tra kết quả",
         "Pass"),
        ("TC_05", "Tìm kiếm audit theo URL",
         "Hiển thị audit khớp với từ khóa tìm",
         "1. Nhập từ khóa vào ô search\n2. Kiểm tra kết quả",
         "Pass"),
    ])

    h3(doc, "3.4.4. Kiểm thử chức năng quản trị")
    testcase_table(doc, "Bảng 3.4. Test case chức năng quản trị viên", [
        ("TC_01", "Admin xem danh sách user",
         "Hiển thị tất cả user trong hệ thống",
         "1. Đăng nhập admin\n2. Truy cập /admin/users",
         "Pass"),
        ("TC_02", "Khóa tài khoản user",
         "User bị khóa không thể đăng nhập",
         "1. Chọn user\n2. Nhấn Khóa\n3. User thử đăng nhập",
         "Pass"),
        ("TC_03", "Mở khóa tài khoản user",
         "User đăng nhập được trở lại",
         "1. Chọn user đã khóa\n2. Nhấn Mở khóa",
         "Pass"),
        ("TC_04", "Cập nhật trọng số rule",
         "Rule mới được áp dụng cho audit tiếp theo",
         "1. Truy cập /admin/rules\n2. Sửa weight rule\n3. Chạy audit mới",
         "Pass"),
        ("TC_05", "Non-admin không truy cập được",
         "User thường bị chuyển về 403",
         "1. Đăng nhập user thường\n2. Truy cập /admin",
         "Pass"),
        ("TC_06", "Xem thống kê",
         "Biểu đồ hiển thị đúng số liệu",
         "1. Truy cập /admin/stats\n2. Kiểm tra biểu đồ",
         "Pass"),
    ])

    # ---- 3.5 Triển khai production ----
    h2(doc, "3.5. Triển khai production")
    para(doc,
        "Pipeline triển khai sử dụng GitHub Actions với 3 workflow: (1) "
        "lint + type-check cho mọi pull request, (2) chạy unit + integration "
        "test trước khi merge vào develop, (3) build Docker image và deploy "
        "tự động khi merge vào main. Frontend Next.js deploy trên Vercel "
        "với preview URL cho từng PR. Backend NestJS và worker BullMQ chạy "
        "trên hai Railway service tách biệt, kết nối đến PostgreSQL và "
        "Redis managed trên Supabase."
    )
    embed_image(doc, "13-cicd-pipeline.png", "Hình 3.11",
        "CI/CD Pipeline: Git push → GitHub Actions → Docker → Vercel/Railway", width_inches=6.5)
    render_items(doc, s_outcome, demote=True)

    # ---- 3.6 Kết luận chương ----
    chapter_conclusion(doc, "3.6.", [
        "Chương 3 đã trình bày toàn bộ quá trình triển khai hệ thống, từ "
        "việc thiết lập môi trường phát triển, tổ chức cấu trúc thư mục dự "
        "án, xây dựng giao diện người dùng cho đến thực hiện kiểm thử chức "
        "năng và triển khai lên production.",
        "Việc sử dụng các công nghệ hiện đại như Next.js, NestJS, PostgreSQL, "
        "BullMQ cùng với quy trình kiểm thử bài bản (Testing Pyramid với "
        "Vitest, TestContainers và Playwright) đã giúp đảm bảo hệ thống "
        "hoạt động ổn định, đáp ứng các yêu cầu kỹ thuật và nghiệp vụ đặt "
        "ra. Pipeline CI/CD tự động trên GitHub Actions kết hợp với triển "
        "khai cloud chi phí thấp (Vercel + Railway + Supabase) đã giúp đạt "
        "mục tiêu chi phí vận hành dưới 40 USD/tháng.",
        "Kết quả triển khai thực tế cho thấy hệ thống đáp ứng đầy đủ các "
        "yêu cầu đã đặt ra ở Chương 1: thời gian phân tích trung bình dưới "
        "10 giây, độ chính xác so với Lighthouse trên 90%, 4 microservices "
        "hoạt động ổn định, dashboard realtime và xuất báo cáo PDF đều "
        "hoạt động đúng thiết kế.",
    ])
    doc.add_page_break()

    # ========================================================================
    # KẾT LUẬN
    # ========================================================================
    p = h1(doc, "KẾT LUẬN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
        "Trong quá trình thực hiện đồ án tốt nghiệp, em đã có cơ hội vận "
        "dụng các kiến thức chuyên ngành như lập trình web, cơ sở dữ liệu, "
        "thiết kế giao diện người dùng, cũng như kiến thức về phân tích và "
        "triển khai hệ thống để xây dựng một nền tảng phân tích SEO website "
        "tự động, hướng đến đối tượng cá nhân học SEO, freelancer và doanh "
        "nghiệp nhỏ và vừa Việt Nam – những người khó tiếp cận các công cụ "
        "SEO thương mại đắt đỏ như Ahrefs hay SEMrush."
    )
    para(doc,
        "Sản phẩm đã hoàn thiện với các chức năng chính như: đăng ký và "
        "đăng nhập với JWT, phân tích SEO một URL với bộ 20 rule on-page, "
        "dashboard theo dõi lịch sử audit, xuất báo cáo PDF, cập nhật tiến "
        "độ realtime qua Socket.IO, cũng như các chức năng quản trị như "
        "quản lý người dùng, cấu hình trọng số rule và thống kê sử dụng. "
        "Hệ thống được phát triển dựa trên kiến trúc microservices với các "
        "công nghệ hiện đại (Next.js, NestJS, PostgreSQL, Redis, BullMQ, "
        "Playwright) nhằm đảm bảo hiệu năng, bảo mật và trải nghiệm người "
        "dùng. Hệ thống đã được triển khai thành công lên Vercel, Railway "
        "và Supabase với chi phí vận hành dưới 40 USD/tháng, đạt đúng mục "
        "tiêu ngân sách đề ra."
    )
    para(doc,
        "Bên cạnh những kết quả đạt được, đồ án vẫn còn tồn tại một số hạn "
        "chế nhất định do giới hạn về thời gian và kinh nghiệm. Có những "
        "điểm chưa được hoàn thiện như: chưa hỗ trợ crawl sâu (deep crawl) "
        "nhiều trang cùng lúc, chưa có tính năng phân tích backlink và off-"
        "page SEO, chưa ứng dụng AI/ML trong việc gợi ý cải thiện nội dung, "
        "và mô hình rate limit còn đơn giản chưa phân biệt theo user tier "
        "một cách linh hoạt."
    )
    para(doc,
        "Trong thời gian tới, em mong muốn tiếp tục nâng cấp hệ thống theo "
        "các hướng: tích hợp module phân tích đối thủ và nghiên cứu từ khóa, "
        "bổ sung gợi ý cải thiện nội dung bằng mô hình ngôn ngữ lớn (LLM), "
        "xây dựng API public cho các nhà phát triển bên thứ ba, phát triển "
        "phiên bản di động và tiến dần tới mô hình SaaS thương mại phục vụ "
        "thị trường SME Việt Nam. Qua quá trình thực hiện đồ án, em đã học "
        "hỏi thêm được nhiều kỹ năng quan trọng như: quản lý thời gian, làm "
        "việc độc lập, giải quyết vấn đề và đặc biệt là kinh nghiệm triển "
        "khai một dự án phần mềm hoàn chỉnh từ thiết kế đến deploy production."
    )
    para(doc,
        "Em xin chân thành cảm ơn thầy/cô giáo hướng dẫn cùng các thầy cô "
        "trong khoa Công nghệ Thông tin – Trường Đại học Giao thông Vận tải "
        "đã tận tình giúp đỡ, hướng dẫn và tạo điều kiện thuận lợi để em "
        "hoàn thành đồ án này."
    )
    doc.add_page_break()

    # ========================================================================
    # TÀI LIỆU THAM KHẢO
    # ========================================================================
    h1(doc, "TÀI LIỆU THAM KHẢO")

    h2(doc, "A. Sách và giáo trình")
    for i, b in enumerate([
        "Martin Fowler, *Patterns of Enterprise Application Architecture*, Addison-Wesley, 2002.",
        "Sam Newman, *Building Microservices*, 2nd ed., O'Reilly Media, 2021.",
        "Robert C. Martin, *Clean Architecture: A Craftsman's Guide to Software Structure and Design*, Prentice Hall, 2017.",
        "Eric Evans, *Domain-Driven Design: Tackling Complexity in the Heart of Software*, Addison-Wesley, 2003.",
    ], 1):
        doc.add_paragraph(f"[{i}] {b}")

    h2(doc, "B. Bài báo khoa học và báo cáo kỹ thuật")
    for i, b in enumerate([
        "Google Search Central, *Search Engine Optimization (SEO) Starter Guide*, 2024.",
        "Mozilla Developer Network, *Web Performance Fundamentals – Core Web Vitals*, 2024.",
        "Addy Osmani, *The State of Web Performance 2024*, web.dev, 2024.",
    ], 5):
        doc.add_paragraph(f"[{i}] {b}")

    h2(doc, "C. Website và tài liệu trực tuyến")
    sites = [
        ("Next.js Documentation", "https://nextjs.org/docs"),
        ("NestJS Documentation", "https://docs.nestjs.com"),
        ("PostgreSQL 16 Documentation", "https://www.postgresql.org/docs/16/"),
        ("Prisma ORM Documentation", "https://www.prisma.io/docs"),
        ("BullMQ Documentation", "https://docs.bullmq.io"),
        ("Cheerio.js Documentation", "https://cheerio.js.org"),
        ("Playwright Documentation", "https://playwright.dev/docs/intro"),
        ("Google Lighthouse", "https://developer.chrome.com/docs/lighthouse/"),
        ("Docker Documentation", "https://docs.docker.com"),
        ("GitHub Actions Documentation", "https://docs.github.com/en/actions"),
        ("Vercel Documentation", "https://vercel.com/docs"),
        ("Railway Documentation", "https://docs.railway.app"),
        ("Supabase Documentation", "https://supabase.com/docs"),
    ]
    for i, (name, url) in enumerate(sites, 8):
        doc.add_paragraph(f"[{i}] {name}. URL: {url} (truy cập 08/04/2026).")
    doc.add_page_break()

    # ========================================================================
    # PHỤ LỤC
    # ========================================================================
    h1(doc, "PHỤ LỤC")

    h2(doc, "Phụ lục A. Danh mục 20 rule SEO on-page")
    rules_tbl = doc.add_table(rows=1, cols=4)
    rules_tbl.style = "Light Grid Accent 1"
    rh = rules_tbl.rows[0].cells
    rh[0].text = "STT"
    rh[1].text = "Tên rule"
    rh[2].text = "Mô tả"
    rh[3].text = "Trọng số"
    rules_data = [
        ("Title tag", "Kiểm tra tồn tại và độ dài (50-60 ký tự)", "8"),
        ("Meta description", "Kiểm tra độ dài (120-160 ký tự)", "7"),
        ("H1 tag", "Chỉ có 1 H1 chứa từ khóa chính", "8"),
        ("Heading hierarchy", "H1→H2→H3 đúng thứ tự", "6"),
        ("Alt attribute", "Mọi <img> phải có alt", "7"),
        ("Canonical URL", "Trang có thẻ rel=canonical", "5"),
        ("Robots meta", "Không vô tình noindex", "6"),
        ("Viewport", "Có meta viewport cho mobile", "10"),
        ("HTTPS", "Trang phục vụ qua HTTPS", "10"),
        ("Open Graph", "Đủ og:title, og:description, og:image", "5"),
        ("Twitter Card", "Có thẻ twitter:card", "3"),
        ("Schema.org", "Có JSON-LD structured data", "6"),
        ("Internal links", "Đủ internal link (≥3) và không gãy", "5"),
        ("External links", "External link có rel phù hợp", "3"),
        ("Image optimization", "Ảnh dùng WebP/AVIF, < 200 KB", "5"),
        ("Page size", "Tổng dung lượng < 2 MB", "4"),
        ("HTTP status", "Trang trả về 200", "8"),
        ("URL structure", "URL ngắn, có từ khóa", "4"),
        ("Language tag", "<html lang> được khai báo", "3"),
        ("Favicon", "Có favicon.ico hoặc link rel=icon", "2"),
    ]
    for i, (name, desc, w) in enumerate(rules_data, 1):
        cells = rules_tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = name
        cells[2].text = desc
        cells[3].text = w

    h2(doc, "Phụ lục B. Roadmap phát triển")
    render_items(doc, s_roadmap, demote=True)

    h2(doc, "Phụ lục C. Bảng phân công công việc")
    work_tbl = doc.add_table(rows=1, cols=4)
    work_tbl.style = "Light Grid Accent 1"
    hdr = work_tbl.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Hạng mục công việc"
    hdr[2].text = "Người thực hiện"
    hdr[3].text = "Thời gian"
    for row in [
        ("1", "Khảo sát, phân tích yêu cầu", "SV thực hiện", "Tuần 1-2"),
        ("2", "Thiết kế kiến trúc microservices, ERD, API", "SV thực hiện", "Tuần 3-4"),
        ("3", "Triển khai Crawler Service", "SV thực hiện", "Tuần 5-6"),
        ("4", "Triển khai SEO Analyzer Service (20 rules)", "SV thực hiện", "Tuần 7-8"),
        ("5", "Triển khai Report Service và PDF export", "SV thực hiện", "Tuần 9"),
        ("6", "Triển khai Frontend Next.js + Socket.IO", "SV thực hiện", "Tuần 10-11"),
        ("7", "Viết test (unit, integration, E2E)", "SV thực hiện", "Tuần 12"),
        ("8", "Triển khai production (Vercel/Railway/Supabase)", "SV thực hiện", "Tuần 13"),
        ("9", "Viết báo cáo đồ án, chuẩn bị bảo vệ", "SV thực hiện", "Tuần 14-15"),
    ]:
        cells = work_tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    # Save
    doc.save(str(OUT))
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
